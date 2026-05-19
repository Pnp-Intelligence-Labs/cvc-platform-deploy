"""
Brambles Strategic Fund DD Pipeline — API Routes
================================================
Prefix: /brambles  (set in main.py)

Endpoints:
  GET    /brambles/companies            — list all companies with status/tier
  POST   /brambles/companies            — add a company from the screening list
  PATCH  /brambles/companies/{id}       — update fields (analyst notes, tier override)
  DELETE /brambles/companies/{id}       — remove a company
  POST   /brambles/companies/{id}/run   — trigger analysis (web-first, no dataroom)
  GET    /brambles/companies/{id}       — get full record including ic_memo_json
"""

from __future__ import annotations

import json
import logging
import os
import re
import sys
import time
from typing import Optional

_BRAMBLES_ENGINE_PATH = "/home/nathan11/repos/brambles_dd_engine"
if _BRAMBLES_ENGINE_PATH not in sys.path:
    sys.path.insert(0, _BRAMBLES_ENGINE_PATH)

import requests
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from pydantic import BaseModel

from api.auth import require_auth
from api.routes.auth import require_jwt, UserInfo
from core.db.connection import get_connection
from llm.openrouter import call as llm_call

logger = logging.getLogger(__name__)

router = APIRouter()

_BRAMBLES_PARTNER_ID = 27
_FULL_ACCESS_ROLES   = {"GP", "Principal", "Director"}


def require_brambles_access(user: UserInfo = Depends(require_jwt)) -> UserInfo:
    """Allow GP/Principal/Director or any PSM assigned to the Brambles partner (id=27)."""
    if user.role in _FULL_ACCESS_ROLES:
        return user
    if _BRAMBLES_PARTNER_ID in (user.assigned_partner_ids or []):
        return user
    raise HTTPException(status_code=403, detail="Access to Brambles DD Platform is restricted")


def _json_safe(obj):
    """Recursively convert Decimal/non-serializable types to JSON-safe equivalents."""
    from decimal import Decimal
    if isinstance(obj, dict):
        return {k: _json_safe(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_json_safe(v) for v in obj]
    if isinstance(obj, Decimal):
        return float(obj)
    return obj


# ---------------------------------------------------------------------------
# MODELS
# ---------------------------------------------------------------------------

class BramblesCompanyInput(BaseModel):
    company_name:     str
    website:          Optional[str] = None
    one_liner:        Optional[str] = None
    employees:        Optional[int] = None
    founded_year:     Optional[int] = None
    hq:               Optional[str] = None
    funding_stage:    Optional[str] = None
    raised_usd_m:     Optional[float] = None
    tech_stack_layer: Optional[str] = None
    relevant_process: Optional[str] = None
    analyst_rationale:Optional[str] = None
    analyst_tier:     Optional[str] = None


class BramblesCompanyUpdate(BaseModel):
    website:          Optional[str] = None
    one_liner:        Optional[str] = None
    employees:        Optional[int] = None
    founded_year:     Optional[int] = None
    hq:               Optional[str] = None
    funding_stage:    Optional[str] = None
    raised_usd_m:     Optional[float] = None
    tech_stack_layer: Optional[str] = None
    relevant_process: Optional[str] = None
    analyst_rationale:Optional[str] = None
    analyst_tier:     Optional[str] = None


# ---------------------------------------------------------------------------
# GET /brambles/companies
# ---------------------------------------------------------------------------

@router.get("/companies")
def list_brambles_companies(user=Depends(require_brambles_access)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT
                    bp.id, bp.company_name, bp.website, bp.one_liner,
                    bp.employees, bp.founded_year, bp.hq,
                    bp.funding_stage, bp.raised_usd_m,
                    bp.tech_stack_layer, bp.relevant_process,
                    bp.analyst_rationale, bp.analyst_tier,
                    bp.status, bp.review_status, bp.tier, bp.tier_label, bp.composite_score,
                    bp.strategic_rationale,
                    bp.pdf_memo_path, bp.pdf_appendix_path, bp.excel_path,
                    bp.review_memo_path,
                    bp.added_by, bp.created_at, bp.updated_at,
                    c.id AS cvc_company_id
                FROM cvc.brambles_pipeline bp
                LEFT JOIN cvc.companies c ON LOWER(c.name) = LOWER(bp.company_name)
                ORDER BY bp.created_at DESC
            """)
            rows = cur.fetchall()

    return [dict(r) for r in rows]


# ---------------------------------------------------------------------------
# GET /brambles/companies/{id}
# ---------------------------------------------------------------------------

@router.get("/companies/{company_id}")
def get_brambles_company(company_id: int, user=Depends(require_brambles_access)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM cvc.brambles_pipeline WHERE id = %s",
                (company_id,)
            )
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")
    result = dict(row)
    # Convert absolute file path → public static URL
    _static_root = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
    rmp = result.get("review_memo_path")
    if rmp and os.path.exists(rmp):
        result["review_memo_url"] = "/static/" + os.path.relpath(rmp, _static_root).replace("\\", "/")
    else:
        result["review_memo_url"] = None
    return result


# ---------------------------------------------------------------------------
# POST /brambles/companies
# ---------------------------------------------------------------------------

@router.post("/companies")
def add_brambles_company(body: BramblesCompanyInput, user=Depends(require_brambles_access)):
    username = user.username if hasattr(user, "username") else "analyst"
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO cvc.brambles_pipeline (
                    company_name, website, one_liner, employees, founded_year,
                    hq, funding_stage, raised_usd_m, tech_stack_layer,
                    relevant_process, analyst_rationale, analyst_tier,
                    status, added_by
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'pending',%s)
                RETURNING id
            """, (
                body.company_name, body.website, body.one_liner, body.employees,
                body.founded_year, body.hq, body.funding_stage, body.raised_usd_m,
                body.tech_stack_layer, body.relevant_process,
                body.analyst_rationale, body.analyst_tier, username,
            ))
            new_id = cur.fetchone()["id"]
            conn.commit()
    return {"created": True, "id": new_id}


# ---------------------------------------------------------------------------
# PATCH /brambles/companies/{id}
# ---------------------------------------------------------------------------

@router.patch("/companies/{company_id}")
def update_brambles_company(company_id: int, body: BramblesCompanyUpdate, user=Depends(require_brambles_access)):
    fields = {k: v for k, v in body.dict().items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    set_clause = ", ".join(f"{k} = %s" for k in fields)
    set_clause += ", updated_at = NOW()"
    values = list(fields.values()) + [company_id]
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE cvc.brambles_pipeline SET {set_clause} WHERE id = %s RETURNING id",
                values
            )
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Company not found")
            conn.commit()
    return {"updated": True}


# ---------------------------------------------------------------------------
# DELETE /brambles/companies/{id}
# ---------------------------------------------------------------------------

@router.delete("/companies/{company_id}", dependencies=[Depends(require_brambles_access)])
def delete_brambles_company(company_id: int):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM cvc.brambles_pipeline WHERE id = %s RETURNING id",
                (company_id,)
            )
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Company not found")
            conn.commit()
    return {"deleted": True}


# ---------------------------------------------------------------------------
# WEB ENRICHMENT HELPERS
# ---------------------------------------------------------------------------

_brave_quota_exhausted: float = 0.0   # unix timestamp when quota was hit; 0 = not exhausted
_BBH_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "8439605362:AAEs0kErefS7YL9JcAx4H_TpTSOSGiBgLrM")
_NATE_CHAT_ID  = os.environ.get("TELEGRAM_CHAT_ID", "8310039682")

def _notify_brave_quota():
    """Fire-and-forget Telegram ping when Brave quota is exhausted."""
    try:
        requests.post(
            f"https://api.telegram.org/bot{_BBH_BOT_TOKEN}/sendMessage",
            json={"chat_id": _NATE_CHAT_ID,
                  "text": "⚠️ Brave Search quota exhausted — Brambles enrichment falling back to list data only. Resets in ~1 hour."},
            timeout=8,
        )
    except Exception:
        pass


def _brave_call(query: str, count: int = 5) -> list:
    """Brave search with primary/backup key failover. Returns raw result list."""
    global _brave_quota_exhausted
    if _brave_quota_exhausted and (time.time() - _brave_quota_exhausted) < 3600:
        return []
    _brave_quota_exhausted = 0.0  # reset if >1 hour has passed
    primary = os.environ.get("BRAVE_SEARCH_KEY", "")
    backup  = os.environ.get("BRAVE_SEARCH_KEY_BACKUP", "")
    for key in filter(None, [primary, backup]):
        try:
            resp = requests.get(
                "https://api.search.brave.com/res/v1/web/search",
                headers={"Accept": "application/json", "X-Subscription-Token": key},
                params={"q": query, "count": count},
                timeout=15,
            )
            resp.raise_for_status()
            return resp.json().get("web", {}).get("results", [])
        except Exception as e:
            err = str(e)
            if any(c in err for c in ("429", "402", "QUOTA")):
                if key == primary and backup:
                    time.sleep(1)
                    continue
                if not _brave_quota_exhausted:
                    import threading
                    threading.Thread(target=_notify_brave_quota, daemon=True).start()
                _brave_quota_exhausted = time.time()
                logger.warning("Brave quota exhausted — enrichment will use list data only.")
                return []
            logger.warning(f"Brave search failed for '{query}': {e}")
            return []
    return []


def _fmt_results(results: list) -> str:
    """Format Brave results into a compact text block for the LLM."""
    if not results:
        return "(no results)"
    lines = []
    for r in results[:5]:
        title   = r.get("title", "")
        url     = r.get("url", "")
        snippet = r.get("description", "") or r.get("extra_snippets", [""])[0] if r.get("extra_snippets") else ""
        lines.append(f"• {title}\n  {url}\n  {snippet[:200]}")
    return "\n".join(lines)


def _enrich_brambles_web(record: dict) -> dict:
    """
    Run 4 targeted Brave searches on a Brambles pipeline company, then
    synthesise the results via LLM into structured scoring signals.

    Returns an enrichment_signals dict. On any failure returns empty signals
    so scoring falls back to company-list data.
    """
    name      = record.get("company_name", "")
    one_liner = record.get("one_liner", "")
    stage     = record.get("funding_stage", "")
    hq        = record.get("hq", "")

    EMPTY = {
        "has_live_deployment": False,
        "deployment_evidence": "",
        "customer_names": [],
        "has_signed_contracts": False,
        "founder_names": [],
        "founder_supply_chain_experience": False,
        "founder_prior_exits": False,
        "founder_background": "",
        "key_facts": [],
        "enrichment_confidence": "none",
        "web_search_ran": False,
    }

    try:
        # ── 4 targeted searches ──────────────────────────────────────────────
        deployment = _brave_call(
            f'"{name}" customers deployment pallet warehouse supply chain pilot 2023 2024 2025 2026',
            count=5,
        )
        funding = _brave_call(
            f'"{name}" funding raised investment round investors announcement',
            count=4,
        )
        founders = _brave_call(
            f'"{name}" founders CEO co-founder background supply chain robotics experience',
            count=4,
        )
        general = _brave_call(
            f'"{name}" {one_liner[:60]} automation robotics news',
            count=4,
        )

        # If all four returned empty, Brave is unavailable — return empty
        if not any([deployment, funding, founders, general]):
            return {**EMPTY, "web_search_ran": False}

        # ── Build numbered source index (for LLM citation) ───────────────────
        all_sources: dict[str, dict] = {}   # id → {title, url, snippet}

        def _index_results(results: list, prefix: str) -> str:
            """Number each result, add to all_sources, return formatted block."""
            lines = []
            for i, r in enumerate(results[:5], 1):
                ref     = f"{prefix}{i}"
                title   = r.get("title", "")
                url     = r.get("url", "")
                snippet = (r.get("description", "") or
                           (r.get("extra_snippets") or [""])[0])[:250]
                if url:
                    all_sources[ref] = {"title": title, "url": url,
                                        "snippet": snippet}
                lines.append(f"[{ref}] {title}\n     {url}\n     {snippet}")
            return "\n".join(lines) if lines else "(no results)"

        d_block = _index_results(deployment, "D")
        f_block = _index_results(funding,    "F")
        r_block = _index_results(founders,   "R")
        g_block = _index_results(general,    "G")

        # ── LLM synthesis with per-claim source citation ──────────────────────
        prompt = f"""You are screening startups for the Brambles Strategic Fund.
Brambles runs global CHEP pallet/container pooling. Relevant companies address
pallet handling, dock automation, warehouse storage, grading/repair, or supply chain intelligence.

Company: {name}
One-liner: {one_liner}
Stage: {stage}  |  HQ: {hq}

WEB RESEARCH — cite sources by their ID (e.g. D1, F2, R3, G1)
--- Deployment & Customers [D1-D5] ---
{d_block}

--- Funding & Investors [F1-F4] ---
{f_block}

--- Founders & Team [R1-R4] ---
{r_block}

--- General / News [G1-G4] ---
{g_block}

Extract signals. Be conservative — only mark True when there is clear evidence above.
For every text field and list, include the source IDs that directly support it.
key_facts: list each as {{"text": "...", "source_ids": ["D1"]}} — max 6 facts, each backed by at least one source.

Output ONLY valid JSON between [JSON_START] and [JSON_END].

[JSON_START]
{{
  "has_live_deployment": false,
  "deployment_evidence": "",
  "deployment_source_ids": [],
  "customer_names": [],
  "customer_source_ids": [],
  "has_signed_contracts": false,
  "founder_names": [],
  "founder_background": "",
  "founder_source_ids": [],
  "founder_supply_chain_experience": false,
  "founder_prior_exits": false,
  "key_facts": [],
  "enrichment_confidence": "low"
}}
[JSON_END]

enrichment_confidence: "low" (few signals), "medium" (some confirmed facts), "high" (multiple verified)
"""

        raw = llm_call(prompt, model="qwen/qwen3-235b-a22b-2507", max_tokens=1500, temperature=0.1)

        m = re.search(r"\[JSON_START\]\s*(\{.*?\})\s*\[JSON_END\]", raw, re.DOTALL)
        if not m:
            m = re.search(r"\{[^{}]*\"has_live_deployment\"[^{}]*\}", raw, re.DOTALL)
        if not m:
            logger.warning(f"Brambles enrichment: could not parse LLM JSON for {name}")
            return {**EMPTY, "web_search_ran": True,
                    "search_sources": [
                        {"query_type": "deployment", "label": "Deployment & Customers",
                         "results": list(all_sources.values())[:5]},
                    ]}

        signals = json.loads(m.group(1) if m.lastindex else m.group(0))

        # ── Resolve source IDs → full URL objects per claim ───────────────────
        def _resolve(ids) -> list:
            return [all_sources[i] for i in (ids or []) if i in all_sources]

        signals["deployment_sources"] = _resolve(signals.pop("deployment_source_ids", []))
        signals["customer_sources"]   = _resolve(signals.pop("customer_source_ids",   []))
        signals["founder_sources"]    = _resolve(signals.pop("founder_source_ids",    []))

        # key_facts: each item is {"text": "...", "source_ids": [...]}
        # Normalise to {"text": "...", "sources": [...]}
        raw_facts = signals.get("key_facts", [])
        normalised_facts = []
        for f in raw_facts:
            if isinstance(f, dict):
                normalised_facts.append({
                    "text":    f.get("text", str(f)),
                    "sources": _resolve(f.get("source_ids", [])),
                })
            else:
                normalised_facts.append({"text": str(f), "sources": []})
        signals["key_facts"] = normalised_facts

        # Keep full raw grouped sources for the side panel
        signals["search_sources"] = [
            {"query_type": "deployment", "label": "Deployment & Customers",
             "results": [v for k, v in all_sources.items() if k.startswith("D")]},
            {"query_type": "funding",    "label": "Funding & Investors",
             "results": [v for k, v in all_sources.items() if k.startswith("F")]},
            {"query_type": "founders",   "label": "Founders & Team",
             "results": [v for k, v in all_sources.items() if k.startswith("R")]},
            {"query_type": "general",    "label": "General / News",
             "results": [v for k, v in all_sources.items() if k.startswith("G")]},
        ]
        signals["web_search_ran"] = True
        return signals

    except Exception as e:
        logger.warning(f"Brambles web enrichment failed for {name}: {e}")
        return {**EMPTY, "web_search_ran": False}


# ---------------------------------------------------------------------------
# POST /brambles/companies/{id}/run  — trigger analysis
# ---------------------------------------------------------------------------

def _run_analysis_bg(company_id: int, record: dict):
    """
    Background task: runs web enrichment + deterministic scoring on a single company row.
    Phase 1 — web enrichment (Brave Search × 4 + LLM synthesis)
    Phase 2 — scoring (FinancialData + ProductData → calculate_brambles_tier)
    Phase 3 — write results to brambles_pipeline
    """
    # Mark as running
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.brambles_pipeline SET status='running', updated_at=NOW() WHERE id=%s",
                (company_id,)
            )
            conn.commit()

    try:
        import sys
        brambles_path = "/home/nathan/repos/brambles_dd_engine"
        if brambles_path not in sys.path:
            sys.path.insert(0, brambles_path)

        from logic.scoring_engine import FinancialData, ProductData, calculate_brambles_tier

        # ── Phase 1: Web enrichment ──────────────────────────────────────────
        enrichment = _enrich_brambles_web(record)

        # ── Phase 2: Build scoring inputs ────────────────────────────────────
        stage = record.get("funding_stage") or "Unknown"
        raised = float(record.get("raised_usd_m") or 0)
        employees = record.get("employees")

        stage_lower = stage.lower()
        is_acquired   = stage_lower == "acquired"
        is_subsidiary = stage_lower in ("parent", "parent entity")
        is_public     = stage_lower == "public"

        # Fold enrichment signals into FinancialData
        has_contracts  = enrichment.get("has_signed_contracts", False)
        customer_names = enrichment.get("customer_names", [])
        customer_count = len(customer_names) if customer_names else None

        fin = FinancialData(
            stage=stage,
            prior_capital_raised_usd=raised * 1_000_000 if raised else None,
            employees=employees,
            is_acquired=is_acquired,
            is_subsidiary=is_subsidiary,
            is_public=is_public,
            has_signed_contracts=has_contracts,
            customer_count=customer_count,
            company_list_mode=True,
        )

        # Map relevant_process to POTF bottleneck
        process_key = (record.get("relevant_process") or "").lower().replace(" ", "_").replace("/", "_").replace("&", "").replace("__", "_").strip("_")
        PROCESS_MAP = {
            "inbound":          "inbound",
            "grading":          "grading",
            "repair":           "repair",
            "storage_outbound": "storage_outbound",
            "storage":          "storage_outbound",
            "outbound":         "storage_outbound",
            "intelligence":     "intelligence",
        }
        bottleneck_fit = PROCESS_MAP.get(process_key)

        # Derive theme alignment from analyst rationale + one_liner + web key facts
        web_facts_text = " ".join(
            f["text"] if isinstance(f, dict) else str(f)
            for f in enrichment.get("key_facts", [])
        ).lower()
        rationale_text = " ".join(filter(None, [
            record.get("analyst_rationale", ""),
            record.get("one_liner", ""),
            record.get("tech_stack_layer", ""),
            web_facts_text,
        ])).lower()

        THEME_KEYWORDS = {
            "freight_decarbonisation": ["decarboni", "emission", "carbon", "freight", "waste valoris"],
            "circularity":             ["circular", "reusable", "pooling", "durable material", "thermoplastic", "composite"],
            "food_supply_chain":       ["food", "cold chain", "shelf life", "refrigerat", "food safety"],
            "digital_transformation":  ["digital transform", "asset track", "rfid", "iot", "s+", "instrument"],
            "e2e_visibility":          ["visibility", "unit load", "end-to-end", "e2e", "optimis", "network optim"],
            "automation_ai":           ["automat", "robot", "ai", "amr", "agv", "autonomous", "warehouse"],
        }
        theme_alignment = [
            k for k, keywords in THEME_KEYWORDS.items()
            if any(kw in rationale_text for kw in keywords)
        ]
        if not theme_alignment and bottleneck_fit:
            PROCESS_THEME = {
                "inbound":          "automation_ai",
                "grading":          "automation_ai",
                "repair":           "automation_ai",
                "storage_outbound": "e2e_visibility",
                "intelligence":     "digital_transformation",
            }
            fallback = PROCESS_THEME.get(bottleneck_fit)
            if fallback:
                theme_alignment = [fallback]

        # Fold enrichment deployment signal into ProductData
        has_deployment = enrichment.get("has_live_deployment", False)

        prd = ProductData(
            bottleneck_fit=bottleneck_fit,
            theme_alignment=theme_alignment,
            has_live_deployment=has_deployment,
        )

        # ── Phase 3: Score ────────────────────────────────────────────────────
        result = calculate_brambles_tier(fin, prd)

        # ── Phase 4: Build memo ───────────────────────────────────────────────
        web_ran = enrichment.get("web_search_ran", False)
        confidence = enrichment.get("enrichment_confidence", "none")
        note = (
            f"Web enrichment ran ({confidence} confidence). "
            f"{len(customer_names)} customer(s) identified. "
            f"Founder research included."
        ) if web_ran else (
            "Generated from company screening list only — web enrichment did not run."
        )

        memo = {
            "company":            record["company_name"],
            "source":             "company_list_web_enriched" if web_ran else "company_list",
            "one_liner":          record.get("one_liner", ""),
            "stage":              stage,
            "raised_usd_m":       raised,
            "hq":                 record.get("hq", ""),
            "employees":          record.get("employees"),
            "founded_year":       record.get("founded_year"),
            "tech_stack_layer":   record.get("tech_stack_layer", ""),
            "relevant_process":   record.get("relevant_process", ""),
            "tier":               result.tier,
            "tier_label":         result.tier_label,
            "composite_score":    result.composite_score,
            "strategic_rationale":result.strategic_rationale,
            "analyst_rationale":  record.get("analyst_rationale", ""),
            "analyst_tier":       record.get("analyst_tier", ""),
            "theme_alignment":    theme_alignment,
            "bottleneck_fit":     bottleneck_fit,
            "rubric_breakdown":   result.rubric_breakdown,
            "flags":              result.flags,
            "fund_fit_detail":    result.fund_fit_detail,
            # Enrichment section
            "enrichment": {
                "web_search_ran":                web_ran,
                "enrichment_confidence":         confidence,
                "has_live_deployment":           enrichment.get("has_live_deployment", False),
                "deployment_evidence":           enrichment.get("deployment_evidence", ""),
                "deployment_sources":            enrichment.get("deployment_sources", []),
                "customer_names":                customer_names,
                "customer_sources":              enrichment.get("customer_sources", []),
                "has_signed_contracts":          has_contracts,
                "founder_names":                 enrichment.get("founder_names", []),
                "founder_supply_chain_experience": enrichment.get("founder_supply_chain_experience", False),
                "founder_prior_exits":           enrichment.get("founder_prior_exits", False),
                "founder_background":            enrichment.get("founder_background", ""),
                "founder_sources":               enrichment.get("founder_sources", []),
                "key_facts":                     enrichment.get("key_facts", []),
            },
            "note": note,
        }

        with get_connection() as conn:
            with conn.cursor() as cur:
                from psycopg2.extras import Json
                cur.execute("""
                    UPDATE cvc.brambles_pipeline SET
                        status           = 'complete',
                        tier             = %s,
                        tier_label       = %s,
                        composite_score  = %s,
                        strategic_rationale = %s,
                        ic_memo_json     = %s,
                        updated_at       = NOW()
                    WHERE id = %s
                """, (
                    result.tier, result.tier_label, result.composite_score,
                    result.strategic_rationale, Json(_json_safe(memo)), company_id,
                ))
                conn.commit()

    except Exception as e:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "UPDATE cvc.brambles_pipeline SET status='failed', updated_at=NOW() WHERE id=%s",
                    (company_id,)
                )
                conn.commit()
        raise


@router.post("/companies/{company_id}/run")
def run_brambles_analysis(
    company_id: int,
    background_tasks: BackgroundTasks,
    user=Depends(require_brambles_access),
):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM cvc.brambles_pipeline WHERE id = %s", (company_id,))
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")
    if row["status"] == "running":
        raise HTTPException(status_code=409, detail="Analysis already running for this company")

    background_tasks.add_task(_run_analysis_bg, company_id, dict(row))
    return {"queued": True, "company": row["company_name"]}


# ---------------------------------------------------------------------------
# POST /brambles/companies/{id}/render  — generate Excel + PDF from ic_memo_json
# ---------------------------------------------------------------------------

_BRAMBLES_OUTPUT_DIR = "/home/nathan11/repos/cvc-intelligence/api/static/brambles"


@router.post("/companies/{company_id}/render")
def render_brambles_docs(company_id: int, user=Depends(require_brambles_access)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM cvc.brambles_pipeline WHERE id = %s", (company_id,))
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")
    if not row["ic_memo_json"]:
        raise HTTPException(status_code=400, detail="Run analysis first — no ic_memo_json available")

    memo = row["ic_memo_json"]
    company_name = row["company_name"]
    safe_name    = re.sub(r"[^\w\-]", "_", company_name)

    out_dir = os.path.join(_BRAMBLES_OUTPUT_DIR, str(company_id))
    os.makedirs(out_dir, exist_ok=True)

    from formatters.report_list_mode import render_xlsx, render_appendix_html_pdf

    xlsx_path     = os.path.join(out_dir, f"{safe_name}_Scorecard.xlsx")
    appendix_path = os.path.join(out_dir, f"{safe_name}_Appendix.html")

    render_xlsx(memo, company_name, xlsx_path)
    render_appendix_html_pdf(memo, company_name, appendix_path)

    # Write paths back to DB
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.brambles_pipeline SET pdf_appendix_path=%s, excel_path=%s, updated_at=NOW() WHERE id=%s",
                (appendix_path, xlsx_path, company_id),
            )
            conn.commit()

    return {
        "xlsx":     f"/brambles/companies/{company_id}/download/xlsx",
        "appendix": f"/brambles/companies/{company_id}/download/appendix",
    }


# ---------------------------------------------------------------------------
# GET /brambles/companies/{id}/download/{type}
# ---------------------------------------------------------------------------

@router.get("/companies/{company_id}/download/{doc_type}")
def download_brambles_doc(company_id: int, doc_type: str, user=Depends(require_brambles_access)):
    from fastapi.responses import FileResponse

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT company_name, excel_path, pdf_appendix_path FROM cvc.brambles_pipeline WHERE id=%s",
                (company_id,),
            )
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")

    if doc_type in ("xlsx", "scorecard"):
        path = row["excel_path"]
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ext   = "xlsx"
    elif doc_type in ("appendix", "memo"):
        # Appendix is HTML (opens in browser with clickable source links)
        path = row["pdf_appendix_path"]
        media = "text/html"
        ext   = "html"
    else:
        raise HTTPException(status_code=400, detail="doc_type must be xlsx/scorecard or appendix/memo")

    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not yet generated — POST /render first")

    safe_name = re.sub(r"[^\w\-]", "_", row["company_name"])
    return FileResponse(path, media_type=media, filename=f"{safe_name}_Brambles.{ext}")


# ---------------------------------------------------------------------------
# POST /brambles/companies/{id}/feedback  — save analyst review verdicts
# GET  /brambles/companies/{id}/feedback  — fetch existing feedback
# ---------------------------------------------------------------------------

class FeedbackItem(BaseModel):
    section:    str
    item_index: int = 0
    item_text:  str
    verdict:    str          # agree | disagree | neutral
    note:       Optional[str] = None
    importance: Optional[int] = None  # 1=Not Important … 5=Critical


class FeedbackPayload(BaseModel):
    items: list[FeedbackItem]


@router.post("/companies/{company_id}/feedback")
def save_brambles_feedback(company_id: int, payload: FeedbackPayload, user=Depends(require_brambles_access)):
    username = user.username if hasattr(user, "username") else "analyst"
    with get_connection() as conn:
        with conn.cursor() as cur:
            for item in payload.items:
                if item.verdict not in ("agree", "disagree", "neutral"):
                    raise HTTPException(status_code=400, detail=f"Invalid verdict: {item.verdict}")
                cur.execute("""
                    INSERT INTO cvc.brambles_feedback
                        (company_id, section, item_index, item_text, verdict, note, importance, reviewed_by, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW())
                    ON CONFLICT (company_id, section, item_index)
                    DO UPDATE SET verdict=EXCLUDED.verdict, note=EXCLUDED.note,
                                  importance=EXCLUDED.importance,
                                  item_text=EXCLUDED.item_text, reviewed_by=EXCLUDED.reviewed_by,
                                  updated_at=NOW()
                """, (company_id, item.section, item.item_index, item.item_text,
                      item.verdict, item.note, item.importance, username))
            conn.commit()
    return {"saved": len(payload.items)}


@router.get("/companies/{company_id}/feedback")
def get_brambles_feedback(company_id: int, user=Depends(require_brambles_access)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT section, item_index, item_text, verdict, note, importance, reviewed_by, updated_at
                FROM cvc.brambles_feedback
                WHERE company_id = %s
                ORDER BY section, item_index
            """, (company_id,))
            rows = cur.fetchall()
    return [dict(r) for r in rows]


# ---------------------------------------------------------------------------
# SECTION WEIGHTS — per-company importance ratings for learning analytics
# ---------------------------------------------------------------------------

class SectionWeightsPayload(BaseModel):
    startup_type: Optional[str] = None
    stage_group:  Optional[str] = None
    weights:      dict  # section -> importance (1–5)


@router.get("/companies/{company_id}/weights")
def get_section_weights(
    company_id:   int,
    startup_type: Optional[str] = None,
    stage_group:  Optional[str] = None,
    user=Depends(require_brambles_access),
):
    with get_connection() as conn:
        with conn.cursor() as cur:
            if startup_type and stage_group:
                cur.execute("""
                    SELECT section, importance, startup_type, stage_group, set_by, updated_at
                    FROM cvc.brambles_section_weights
                    WHERE pipeline_id = %s AND startup_type = %s AND stage_group = %s
                    ORDER BY section
                """, (company_id, startup_type, stage_group))
            else:
                cur.execute("""
                    SELECT section, importance, startup_type, stage_group, set_by, updated_at
                    FROM cvc.brambles_section_weights
                    WHERE pipeline_id = %s
                    ORDER BY section
                """, (company_id,))
            rows = cur.fetchall()
    return [dict(r) for r in rows]


@router.post("/companies/{company_id}/weights")
def save_section_weights(company_id: int, body: SectionWeightsPayload, user=Depends(require_brambles_access)):
    username = user.username if hasattr(user, "username") else "analyst"
    with get_connection() as conn:
        with conn.cursor() as cur:
            for section, importance in body.weights.items():
                if not (1 <= int(importance) <= 5):
                    raise HTTPException(400, f"Importance must be 1–5, got {importance} for {section}")
                cur.execute("""
                    INSERT INTO cvc.brambles_section_weights
                        (pipeline_id, startup_type, stage_group, section, importance, set_by, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, NOW())
                    ON CONFLICT (pipeline_id, startup_type, stage_group, section) DO UPDATE SET
                        importance   = EXCLUDED.importance,
                        set_by       = EXCLUDED.set_by,
                        updated_at   = NOW()
                """, (company_id, body.startup_type, body.stage_group, section, int(importance), username))
        conn.commit()
    return {"saved": len(body.weights)}


# ---------------------------------------------------------------------------
# REVIEW MEMO — background generation
# ---------------------------------------------------------------------------

def _render_review_memo_html(review_memo: dict, ic_memo: dict, company_name: str, path: str):
    """Render the SLAM Intelligence review memo as a professional, printable HTML document."""
    from datetime import date

    cvc_vs    = review_memo.get("cvc_vs_brambles") or review_memo.get("cvc_vs_bramles", {})
    agreement = cvc_vs.get("agreement", "")
    agree_color = {"agree": "#1B6F3A", "partial": "#C8701A", "disagree": "#B91C1C"}.get(agreement, "#475569")
    agree_label = {"agree": "Aligned with Brambles", "partial": "Partially Aligned", "disagree": "Differs from Brambles"}.get(agreement, agreement.title())

    rec       = review_memo.get("recommendation", "—")
    rec_color = {"Pursue": "#1B6F3A", "Monitor": "#C8701A", "Pass": "#B91C1C"}.get(rec, "#1B3A4B")

    def p(text):
        return f"<p>{text}</p>" if text and text != "—" else ""

    def ul(items):
        if not items:
            return "<p class='empty'>None identified.</p>"
        return "<ul>" + "".join(f"<li>{i}</li>" for i in items) + "</ul>"

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>SLAM Intelligence — Investment Memo — {company_name}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Merriweather:ital,wght@0,400;0,700;1,400&family=Inter:wght@400;500;600;700&display=swap');
  *, *::before, *::after {{ box-sizing: border-box; }}
  body {{
    font-family: 'Merriweather', Georgia, serif;
    font-size: 14px;
    line-height: 1.8;
    color: #1a1a1a;
    background: #fff;
    max-width: 820px;
    margin: 0 auto;
    padding: 56px 64px;
  }}
  /* Header */
  .doc-header {{
    border-bottom: 3px solid #1B3A4B;
    padding-bottom: 20px;
    margin-bottom: 8px;
  }}
  .doc-header .firm {{ font-family: 'Inter', sans-serif; font-size: 11px; font-weight: 700;
    letter-spacing: .12em; text-transform: uppercase; color: #C8972B; margin-bottom: 10px; }}
  .doc-header h1 {{ font-family: 'Inter', sans-serif; font-size: 22px; font-weight: 700;
    color: #1B3A4B; margin: 0 0 6px; line-height: 1.3; }}
  .doc-meta {{ font-family: 'Inter', sans-serif; font-size: 12px; color: #64748b;
    display: flex; gap: 20px; flex-wrap: wrap; margin-top: 8px; }}
  .doc-meta span {{ white-space: nowrap; }}
  .confidential {{ font-family: 'Inter', sans-serif; font-size: 10px; font-weight: 600;
    letter-spacing: .1em; text-transform: uppercase; color: #94a3b8; margin-top: 6px; }}
  /* Sections */
  .section {{ margin-top: 36px; }}
  .section-label {{ font-family: 'Inter', sans-serif; font-size: 10px; font-weight: 700;
    letter-spacing: .1em; text-transform: uppercase; color: #94a3b8;
    border-top: 1px solid #e2e8f0; padding-top: 10px; margin-bottom: 10px; }}
  p {{ margin: 0 0 14px; }}
  p.empty {{ color: #94a3b8; font-style: italic; margin: 0; }}
  ul {{ padding-left: 20px; margin: 0 0 14px; }}
  li {{ margin-bottom: 7px; }}
  /* Tier comparison box */
  .tier-compare {{
    display: flex; align-items: center; gap: 0;
    border: 1px solid #e2e8f0; border-radius: 8px;
    overflow: hidden; margin: 16px 0;
    font-family: 'Inter', sans-serif;
  }}
  .tier-cell {{
    flex: 1; padding: 14px 20px; text-align: center;
    border-right: 1px solid #e2e8f0;
  }}
  .tier-cell:last-child {{ border-right: none; }}
  .tier-cell .tc-label {{ font-size: 10px; font-weight: 600; letter-spacing:.08em;
    text-transform:uppercase; color: #94a3b8; margin-bottom: 4px; }}
  .tier-cell .tc-val {{ font-size: 18px; font-weight: 700; color: #1B3A4B; }}
  .tier-cell .tc-sub {{ font-size: 11px; color: #64748b; margin-top: 2px; }}
  .tier-verdict {{
    padding: 14px 24px; text-align: center; min-width: 160px;
    background: {agree_color}; color: white;
  }}
  .tier-verdict .tv-label {{ font-size: 10px; font-weight: 600; letter-spacing:.08em;
    text-transform:uppercase; opacity:.8; margin-bottom: 4px; }}
  .tier-verdict .tv-val {{ font-size: 13px; font-weight: 700; }}
  /* Rationale block */
  .rationale-block {{
    background: #f8fafc;
    border-left: 4px solid {agree_color};
    padding: 14px 18px;
    border-radius: 0 6px 6px 0;
    font-size: 13.5px;
    line-height: 1.75;
    margin-bottom: 16px;
  }}
  /* Recommendation */
  .rec-block {{
    display: flex; align-items: flex-start; gap: 16px;
    padding: 20px 24px; border-radius: 8px;
    border: 2px solid {rec_color};
    margin-bottom: 16px;
  }}
  .rec-badge {{
    font-family: 'Inter', sans-serif;
    font-size: 13px; font-weight: 700;
    color: white; background: {rec_color};
    padding: 6px 16px; border-radius: 4px;
    white-space: nowrap; margin-top: 1px;
  }}
  .rec-text {{ font-size: 13.5px; line-height: 1.7; }}
  /* Score bar */
  .score-row {{ font-family:'Inter',sans-serif; font-size:12px; color:#475569;
    display:flex; align-items:center; gap:10px; margin-bottom:6px; }}
  .score-bar-bg {{ flex:1; height:6px; background:#e2e8f0; border-radius:3px; }}
  .score-bar-fill {{ height:6px; border-radius:3px; background:#1B3A4B; }}
  /* Footer */
  .doc-footer {{
    margin-top: 56px; padding-top: 16px;
    border-top: 1px solid #e2e8f0;
    font-family: 'Inter', sans-serif;
    font-size: 10px; color: #94a3b8;
    display: flex; justify-content: space-between;
  }}
  @media print {{
    body {{ padding: 32px 48px; }}
    .tier-compare, .rec-block, .rationale-block {{ break-inside: avoid; }}
  }}
</style>
</head>
<body>

<div class="doc-header">
  <div class="firm">SLAM Intelligence</div>
  <h1>Investment Assessment Memo<br>{company_name}</h1>
  <div class="doc-meta">
    <span><strong>Stage:</strong> {ic_memo.get('stage','—')}</span>
    <span><strong>Raised:</strong> ${ic_memo.get('raised_usd_m') or '—'}M</span>
    <span><strong>HQ:</strong> {ic_memo.get('hq','—')}</span>
    <span><strong>Employees:</strong> {ic_memo.get('employees') or '—'}</span>
    <span><strong>Score:</strong> {ic_memo.get('composite_score','—')}/100</span>
    <span><strong>Date:</strong> {date.today().strftime('%B %d, %Y')}</span>
  </div>
  <div class="confidential">Confidential — Brambles Strategic Fund Advisory</div>
</div>

<div class="section">
  <div class="section-label">Investment Overview</div>
  {p(review_memo.get('investment_overview') or review_memo.get('executive_summary',''))}
</div>

<div class="section">
  <div class="section-label">SLAM Intelligence Position vs. Brambles Assessment</div>
  <div class="tier-compare">
    <div class="tier-cell">
      <div class="tc-label">Brambles Assessment</div>
      <div class="tc-val">{cvc_vs.get('brambles_tier','—')}</div>
      <div class="tc-sub">{cvc_vs.get('brambles_tier_label','')}</div>
    </div>
    <div class="tier-cell">
      <div class="tc-label">SLAM Independent Score</div>
      <div class="tc-val">{cvc_vs.get('cvc_tier','—')}</div>
      <div class="tc-sub">{ic_memo.get('tier_label','')}</div>
    </div>
    <div class="tier-verdict">
      <div class="tv-label">SLAM Verdict</div>
      <div class="tv-val">{agree_label}</div>
    </div>
  </div>
  <div class="rationale-block">{cvc_vs.get('rationale','—')}</div>
</div>

<div class="section">
  <div class="section-label">Commercial Traction</div>
  {p(review_memo.get('commercial_traction',''))}
</div>

<div class="section">
  <div class="section-label">Founding Team</div>
  {p(review_memo.get('founding_team',''))}
</div>

<div class="section">
  <div class="section-label">Key Evidence</div>
  {ul(review_memo.get('key_evidence', review_memo.get('confirmed_strengths', [])))}
</div>

<div class="section">
  <div class="section-label">Risks &amp; Open Questions</div>
  {p(review_memo.get('risks_and_gaps') or '') or ul(review_memo.get('concerns', []))}
</div>

<div class="section">
  <div class="section-label">Brambles Strategic Fit</div>
  {p(review_memo.get('brambles_fit',''))}
</div>

<div class="section">
  <div class="section-label">SLAM Recommendation</div>
  <div class="rec-block">
    <div class="rec-badge">{rec}</div>
    <div class="rec-text">{review_memo.get('recommendation_paragraph') or review_memo.get('recommendation_rationale','')}</div>
  </div>
</div>

<div class="doc-footer">
  <span>SLAM Intelligence — Independent Advisory</span>
  <span>Brambles Strategic Fund</span>
  <span>{date.today().strftime('%Y')}</span>
</div>

</body>
</html>"""

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)


def _update_brambles_learning():
    """
    Aggregate finalized analyst feedback patterns across all companies and store
    as 'brambles_agent_learning' in platform_settings. Loaded by future enrichment
    runs to calibrate source skepticism.
    """
    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT bf.section, bf.verdict, bf.item_text, bf.note,
                           bp.company_name
                    FROM cvc.brambles_feedback bf
                    JOIN cvc.brambles_pipeline bp ON bp.id = bf.company_id
                    WHERE bp.review_status = 'finalized'
                    ORDER BY bf.section, bf.verdict
                """)
                rows = cur.fetchall()

        if not rows:
            return

        from collections import Counter
        section_verdicts: dict = {}
        for r in rows:
            sec = r["section"]
            v   = r["verdict"]
            section_verdicts.setdefault(sec, []).append(v)

        patterns = {}
        for sec, verdicts in section_verdicts.items():
            c = Counter(verdicts)
            total = len(verdicts)
            patterns[sec] = {
                "agree_pct":    round(c.get("agree",    0) / total * 100),
                "disagree_pct": round(c.get("disagree", 0) / total * 100),
                "neutral_pct":  round(c.get("neutral",  0) / total * 100),
                "total":        total,
            }

        # Most-disputed sections (disagree > 40%)
        flags = [s for s, p in patterns.items() if p["disagree_pct"] >= 40]

        learning = {
            "section_accuracy": patterns,
            "high_dispute_sections": flags,
            "note": (
                f"Based on {len(rows)} verdicts across "
                f"{len({r['company_name'] for r in rows})} reviewed companies. "
                f"Sections with >40% disagreement: {', '.join(flags) or 'none'}."
            ),
        }

        with get_connection() as conn:
            with conn.cursor() as cur:
                from psycopg2.extras import Json
                cur.execute("""
                    INSERT INTO cvc.platform_settings (key, value)
                    VALUES ('brambles_agent_learning', %s)
                    ON CONFLICT (key) DO UPDATE
                      SET value = EXCLUDED.value
                """, (json.dumps(learning),))
                conn.commit()

        logger.info(f"Brambles learning updated: {learning['note']}")
    except Exception as e:
        logger.warning(f"Brambles learning update failed: {e}")


def _generate_review_memo_bg(company_id: int):
    """
    Background task: synthesise a CVC vs Brambles review memo using analyst verdicts
    and the ic_memo_json. Stores result in review_memo_json + renders HTML.
    """
    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT * FROM cvc.brambles_pipeline WHERE id = %s", (company_id,))
                row = cur.fetchone()
                cur.execute("""
                    SELECT section, item_index, item_text, verdict, note
                    FROM cvc.brambles_feedback WHERE company_id = %s
                    ORDER BY section, item_index
                """, (company_id,))
                feedback_rows = cur.fetchall()

        memo             = row["ic_memo_json"] or {}
        enr              = memo.get("enrichment", {})
        brambles_tier    = row.get("analyst_tier") or "Unknown"
        cvc_tier         = row.get("tier")
        cvc_tier_label   = row.get("tier_label") or ""
        company_name     = row["company_name"]

        agreed    = [r for r in feedback_rows if r["verdict"] == "agree"]
        disagreed = [r for r in feedback_rows if r["verdict"] == "disagree"]
        neutral   = [r for r in feedback_rows if r["verdict"] == "neutral"]

        def fmt_rows(rows):
            if not rows:
                return "(none)"
            lines = []
            for r in rows:
                note = (r.get("note") or "").strip()
                if note:
                    lines.append(f"- {r['item_text']}\n  → Analyst reasoning: \"{note}\"")
                else:
                    lines.append(f"- {r['item_text']}")
            return "\n".join(lines)

        key_facts_text = "\n".join(
            f"- {f['text'] if isinstance(f, dict) else f}"
            for f in enr.get("key_facts", [])
        )

        prompt = f"""You are a senior investment analyst at SLAM Intelligence, writing an
independent third-party investment memo for the Brambles Strategic Fund.

CONTEXT:
Brambles operates global CHEP pallet and container pooling — roughly 330 million pallets across
~60 countries. Their automation challenge is brownfield: no rip-and-replace, must work alongside
existing workflows, needs global scalability. That lens should inform everything you write.

────────────────────────────────────────
COMPANY DATA
────────────────────────────────────────
Company:        {company_name}
One-liner:      {memo.get('one_liner','')}
Stage:          {memo.get('stage','')}
HQ:             {memo.get('hq','')}
Raised:         ${memo.get('raised_usd_m') or '?'}M
Employees:      {memo.get('employees') or '?'}
Founded:        {memo.get('founded_year') or '?'}
Bottleneck fit: {memo.get('bottleneck_fit','').replace('_',' ')}
Themes:         {', '.join(memo.get('theme_alignment', []))}
CVC Score:      {memo.get('composite_score','?')}/100

BRAMBLES' ASSESSMENT:         {brambles_tier}
SLAM INTELLIGENCE SCORING:    Tier {cvc_tier} — {cvc_tier_label}

────────────────────────────────────────
ANALYST REVIEW (from page-by-page verification)
────────────────────────────────────────
CONFIRMED (analyst agreed — use as evidence):
{fmt_rows(agreed)}

DISPUTED (analyst disagreed — include reasoning as a concern or caveat):
{fmt_rows(disagreed)}

UNVERIFIED (analyst marked neutral — mention only with appropriate hedging):
{fmt_rows(neutral)}

WEB ENRICHMENT:
- Deployment confirmed: {enr.get('has_live_deployment')}
- Deployment evidence: {enr.get('deployment_evidence') or 'none found'}
- Named customers: {', '.join(enr.get('customer_names', [])) or 'none found'}
- Signed contracts: {enr.get('has_signed_contracts')}
- Founder background: {enr.get('founder_background') or 'not found'}
- SC experience: {enr.get('founder_supply_chain_experience')} | Prior exits: {enr.get('founder_prior_exits')}
- Enrichment confidence: {enr.get('enrichment_confidence')}

KEY FACTS FROM RESEARCH:
{key_facts_text or '(none)'}

SCORING FLAGS: {'; '.join(memo.get('flags', [])) or 'none'}

────────────────────────────────────────
WRITING INSTRUCTIONS
────────────────────────────────────────
Write a professional investment memo — the kind a senior analyst at a top-tier fund
would hand to a client. The tone is direct, evidence-based, and opinionated.

Rules:
- Every section is flowing PROSE — not bullets, not a list of facts.
- Weave the analyst's confirmed notes and disagreement reasoning naturally into the
  relevant section. Do NOT say "the analyst noted". Just use the reasoning as your own.
- Disputed claims should appear as concerns or caveats in the relevant section,
  with the analyst's reasoning as your explanation for why.
- Unverified items should be mentioned only with hedging ("reportedly", "per company claims").
- The "cvc_vs_brambles.rationale" paragraph is the centrepiece — make it specific and bold.
  If you agree with Brambles, explain precisely what evidence confirms their view.
  If you differ, say exactly what they missed, overstated, or underestimated.
- Always refer to the authoring firm as "SLAM Intelligence" — never "CVC" or "Claw Venture Capital".
- "commercial_traction" and "founding_team" should each be a full paragraph.
- "risks_and_gaps" should be a paragraph, not a list.
- "recommendation_paragraph" should read as a closing statement with a clear next step.
- key_evidence: 4–6 short bullet strings of the most important confirmed facts only.

Output ONLY valid JSON between [JSON_START] and [JSON_END]:

[JSON_START]
{{
  "investment_overview": "2-3 sentence paragraph: what the company does, why it matters to Brambles, and SLAM Intelligence's overall stance",
  "cvc_vs_brambles": {{
    "brambles_tier": "{brambles_tier}",
    "brambles_tier_label": "",
    "cvc_tier": "Tier {cvc_tier}",
    "agreement": "agree",
    "rationale": "Full paragraph — SLAM Intelligence's independent verdict on this company vs Brambles' assessment. Specific, evidence-driven, opinionated."
  }},
  "commercial_traction": "Full paragraph on deployment status, named customers, contract evidence — incorporating analyst verdicts naturally",
  "founding_team": "Full paragraph on the founding team, their relevant background, and SLAM Intelligence's read on team quality",
  "key_evidence": ["4-6 confirmed factual bullets supporting the thesis"],
  "risks_and_gaps": "Full paragraph on concerns, disputed claims, open questions — written as SLAM Intelligence's own analysis",
  "brambles_fit": "Full paragraph on how specifically this maps to CHEP operations — brownfield fit, target bottleneck, global scalability",
  "recommendation": "Pursue",
  "recommendation_paragraph": "Closing paragraph — SLAM Intelligence recommendation with clear rationale and suggested next step"
}}
[JSON_END]

recommendation must be exactly one of: Pursue / Monitor / Pass
agreement must be exactly one of: agree / partial / disagree"""

        raw = llm_call(prompt, model="qwen/qwen3-235b-a22b-2507", max_tokens=2000, temperature=0.2)

        m = re.search(r"\[JSON_START\]\s*(\{.*?\})\s*\[JSON_END\]", raw, re.DOTALL)
        if not m:
            m = re.search(r"\{[^{}]*\"executive_summary\"[^{}]*\}", raw, re.DOTALL)
        if not m:
            logger.warning(f"Review memo: could not parse LLM JSON for company {company_id}")
            return

        review_memo = json.loads(m.group(1) if m.lastindex else m.group(0))

        # Scrub any residual firm-name leakage from LLM output
        def _scrub(obj):
            if isinstance(obj, str):
                return (obj
                    .replace("Claw Venture Capital", "SLAM Intelligence")
                    .replace("CVC recommends", "SLAM Intelligence recommends")
                    .replace("CVC views", "SLAM Intelligence views")
                    .replace("CVC rates", "SLAM Intelligence rates")
                    .replace("CVC's", "SLAM Intelligence's")
                    .replace(" CVC ", " SLAM Intelligence "))
            if isinstance(obj, dict):
                return {k: _scrub(v) for k, v in obj.items()}
            if isinstance(obj, list):
                return [_scrub(v) for v in obj]
            return obj
        review_memo = _scrub(review_memo)

        # Render HTML
        safe_name = re.sub(r"[^\w\-]", "_", company_name)
        out_dir   = os.path.join(_BRAMBLES_OUTPUT_DIR, str(company_id))
        os.makedirs(out_dir, exist_ok=True)
        memo_path = os.path.join(out_dir, f"{safe_name}_Review_Memo.html")
        _render_review_memo_html(review_memo, memo, company_name, memo_path)

        with get_connection() as conn:
            with conn.cursor() as cur:
                from psycopg2.extras import Json
                cur.execute(
                    """UPDATE cvc.brambles_pipeline
                       SET review_memo_json=%s, review_memo_path=%s, updated_at=NOW()
                       WHERE id=%s""",
                    (Json(_json_safe(review_memo)), memo_path, company_id)
                )
                conn.commit()

        logger.info(f"Review memo generated for company {company_id}")
        _update_brambles_learning()

    except Exception as e:
        logger.warning(f"Review memo generation failed for company {company_id}: {e}")
        try:
            with get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE cvc.brambles_pipeline SET review_status='generation_failed', updated_at=NOW() WHERE id=%s",
                        (company_id,)
                    )
                    conn.commit()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# POST /brambles/companies/{id}/finish-review
# ---------------------------------------------------------------------------

@router.post("/companies/{company_id}/finish-review")
def finish_brambles_review(
    company_id: int,
    background_tasks: BackgroundTasks,
    user=Depends(require_brambles_access),
):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT ic_memo_json, review_status FROM cvc.brambles_pipeline WHERE id = %s", (company_id,))
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")
    if not row["ic_memo_json"]:
        raise HTTPException(status_code=400, detail="Run analysis first — no ic_memo_json available")

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.brambles_pipeline SET review_status='finalized', updated_at=NOW() WHERE id=%s",
                (company_id,)
            )
            conn.commit()

    background_tasks.add_task(_generate_review_memo_bg, company_id)
    return {"finalized": True, "memo_generating": True}


# ---------------------------------------------------------------------------
# GET /brambles/companies/{id}/download/review-memo
# ---------------------------------------------------------------------------

@router.get("/companies/{company_id}/download/review-memo")
def download_review_memo(company_id: int, user=Depends(require_brambles_access)):
    from fastapi.responses import FileResponse
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT company_name, review_memo_path FROM cvc.brambles_pipeline WHERE id=%s",
                (company_id,),
            )
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Company not found")
    if not row["review_memo_path"] or not os.path.exists(row["review_memo_path"]):
        raise HTTPException(status_code=404, detail="Review memo not yet generated")
    safe_name = re.sub(r"[^\w\-]", "_", row["company_name"])
    return FileResponse(row["review_memo_path"], media_type="text/html",
                        filename=f"{safe_name}_SLAM_Review_Memo.html")


# ---------------------------------------------------------------------------
# POST /brambles/companies/{id}/reopen-review
# ---------------------------------------------------------------------------

@router.post("/companies/{company_id}/reopen-review")
def reopen_brambles_review(company_id: int, user=Depends(require_brambles_access)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM cvc.brambles_pipeline WHERE id = %s", (company_id,))
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Company not found")
            cur.execute(
                """UPDATE cvc.brambles_pipeline
                   SET review_status='pending', review_memo_json=NULL,
                       review_memo_path=NULL, updated_at=NOW()
                   WHERE id=%s""",
                (company_id,)
            )
            conn.commit()
    return {"reopened": True}


# ---------------------------------------------------------------------------
# GET /brambles/companies/{id}/download/review-memo.pdf
# GET /brambles/companies/{id}/download/review-memo.docx
# ---------------------------------------------------------------------------

def _get_memo_row(company_id: int):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT company_name, review_memo_path, review_memo_json, ic_memo_json FROM cvc.brambles_pipeline WHERE id=%s",
                (company_id,),
            )
            return cur.fetchone()


@router.get("/companies/{company_id}/memo-pdf")
def download_review_memo_pdf(company_id: int, user=Depends(require_brambles_access)):
    from fastapi.responses import Response
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    import io
    from datetime import date

    row = _get_memo_row(company_id)
    if not row:
        raise HTTPException(404, "Company not found")
    memo = row["review_memo_json"]
    ic   = row["ic_memo_json"] or {}
    if not memo:
        raise HTTPException(404, "Review memo not yet generated")

    cvc_vs    = memo.get("cvc_vs_brambles") or memo.get("cvc_vs_bramles", {})
    agreement = cvc_vs.get("agreement", "")
    rec       = memo.get("recommendation", "—")
    company_name = row["company_name"]
    safe_name = re.sub(r"[^\w\-]", "_", company_name)

    SLATE  = colors.HexColor("#1B3A4B")
    GOLD   = colors.HexColor("#C8972B")
    MUTED  = colors.HexColor("#64748b")
    LIGHT  = colors.HexColor("#94a3b8")
    REC_COLORS = {"Pursue": "#1B6F3A", "Monitor": "#C8701A", "Pass": "#B91C1C"}
    rec_hex = colors.HexColor(REC_COLORS.get(rec, "#1B3A4B"))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=1.1*inch, rightMargin=1.1*inch,
                            topMargin=1*inch, bottomMargin=1*inch)

    styles = getSampleStyleSheet()
    body_style  = ParagraphStyle("body",  fontSize=10, leading=16, textColor=colors.HexColor("#1a1a1a"), spaceAfter=10)
    label_style = ParagraphStyle("label", fontSize=7.5, leading=10, textColor=LIGHT, fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=4, wordWrap='LTR')
    h1_style    = ParagraphStyle("h1",    fontSize=18, leading=22, textColor=SLATE, fontName="Helvetica-Bold", spaceAfter=4)
    firm_style  = ParagraphStyle("firm",  fontSize=9,  leading=12, textColor=GOLD,  fontName="Helvetica-Bold", spaceAfter=6)
    meta_style  = ParagraphStyle("meta",  fontSize=8.5, leading=12, textColor=MUTED, spaceAfter=4)
    bullet_style= ParagraphStyle("bullet",fontSize=10, leading=15, textColor=colors.HexColor("#1a1a1a"), leftIndent=14, spaceAfter=5)

    story = []

    # Header
    story.append(Paragraph("SLAM INTELLIGENCE", firm_style))
    story.append(Paragraph(f"Investment Assessment Memo — {company_name}", h1_style))
    meta_parts = [
        f"Stage: {ic.get('stage','—')}",
        f"Raised: ${ic.get('raised_usd_m') or '—'}M",
        f"Score: {ic.get('composite_score','—')}/100",
        f"Date: {date.today().strftime('%B %d, %Y')}",
    ]
    story.append(Paragraph("  ·  ".join(meta_parts), meta_style))
    story.append(Paragraph("Confidential — Brambles Strategic Fund Advisory", meta_style))
    story.append(HRFlowable(width="100%", thickness=2, color=SLATE, spaceAfter=14))

    # Tier comparison table
    story.append(Paragraph("BRAMBLES VS. SLAM ASSESSMENT", label_style))
    agree_label_map = {"agree": "Aligned", "partial": "Partial", "disagree": "Differs"}
    tier_data = [
        ["Brambles Assessment", "SLAM Score", "Verdict"],
        [cvc_vs.get("brambles_tier","—"), cvc_vs.get("cvc_tier","—"), agree_label_map.get(agreement, agreement.title())],
    ]
    tier_table = Table(tier_data, colWidths=[2.2*inch, 2.2*inch, 2.2*inch])
    tier_table.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0), SLATE),
        ("TEXTCOLOR",    (0,0), (-1,0), colors.white),
        ("FONTNAME",     (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 9),
        ("ALIGN",        (0,0), (-1,-1), "CENTER"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#f8fafc")]),
        ("GRID",         (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
        ("TOPPADDING",   (0,0), (-1,-1), 8),
        ("BOTTOMPADDING",(0,0), (-1,-1), 8),
    ]))
    story.append(tier_table)
    story.append(Spacer(1, 8))

    # Body sections
    sections_data = [
        ("INVESTMENT OVERVIEW",            memo.get("investment_overview", "")),
        ("SLAM INTELLIGENCE VS. BRAMBLES", cvc_vs.get("rationale", "")),
        ("COMMERCIAL TRACTION",            memo.get("commercial_traction", "")),
        ("FOUNDING TEAM",                  memo.get("founding_team", "")),
        ("RISKS & OPEN QUESTIONS",         memo.get("risks_and_gaps", "")),
        ("BRAMBLES STRATEGIC FIT",         memo.get("brambles_fit", "")),
    ]
    for label, body in sections_data:
        if body:
            story.append(Paragraph(label, label_style))
            story.append(Paragraph(body, body_style))

    # Key evidence
    if memo.get("key_evidence"):
        story.append(Paragraph("KEY EVIDENCE", label_style))
        for item in memo["key_evidence"]:
            story.append(Paragraph(f"• {item}", bullet_style))
        story.append(Spacer(1, 6))

    # Recommendation
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=10, spaceAfter=10))
    story.append(Paragraph("SLAM RECOMMENDATION", label_style))
    rec_style = ParagraphStyle("rec", fontSize=11, leading=16,
                               textColor=colors.HexColor("#1a1a1a"), spaceAfter=6)
    rec_label = ParagraphStyle("reclabel", fontSize=13, leading=18,
                               textColor=rec_hex, fontName="Helvetica-Bold", spaceAfter=4)
    story.append(Paragraph(rec, rec_label))
    story.append(Paragraph(memo.get("recommendation_paragraph", ""), rec_style))

    doc.build(story)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_SLAM_Review_Memo.pdf"'},
    )


@router.get("/companies/{company_id}/memo-docx")
def download_review_memo_docx(company_id: int, user=Depends(require_brambles_access)):
    from fastapi.responses import Response
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    row = _get_memo_row(company_id)
    if not row:
        raise HTTPException(404, "Company not found")
    memo = row["review_memo_json"]
    ic   = row["ic_memo_json"] or {}
    if not memo:
        raise HTTPException(404, "Review memo not yet generated")

    cvc_vs    = memo.get("cvc_vs_brambles") or memo.get("cvc_vs_bramles", {})
    agreement = cvc_vs.get("agreement", "")
    rec       = memo.get("recommendation", "—")
    from datetime import date

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def add_heading(text, level=1, color=(27, 58, 75)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_para(text, italic=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text or "")
        run.font.size = Pt(size)
        run.italic = italic
        return p

    def add_label(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(148, 163, 184)
        return p

    # Header
    title = doc.add_heading("Investment Assessment Memo", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(27, 58, 75)
    p = doc.add_paragraph()
    run = p.add_run("SLAM Intelligence  ·  Brambles Strategic Fund  ·  Confidential")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(200, 151, 43)
    run.font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(row["company_name"])
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 58, 75)

    meta_items = [
        f"Stage: {ic.get('stage','—')}",
        f"Raised: ${ic.get('raised_usd_m') or '—'}M",
        f"Score: {ic.get('composite_score','—')}/100",
        f"Date: {date.today().strftime('%B %d, %Y')}",
    ]
    add_para("  ·  ".join(meta_items), size=9).runs[0].font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    # Tier comparison table
    add_label("Brambles vs. SLAM Assessment")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = f"Brambles\n{cvc_vs.get('brambles_tier','—')}"
    hdr[1].text = f"SLAM Score\n{cvc_vs.get('cvc_tier','—')}"
    verdict_map = {"agree": "Aligned", "partial": "Partial", "disagree": "Differs"}
    hdr[2].text = f"Verdict\n{verdict_map.get(agreement, agreement.title())}"
    doc.add_paragraph()

    # Body sections
    sections_data = [
        ("Investment Overview",           memo.get("investment_overview", "")),
        ("SLAM Intelligence vs. Brambles", cvc_vs.get("rationale", "")),
        ("Commercial Traction",            memo.get("commercial_traction", "")),
        ("Founding Team",                  memo.get("founding_team", "")),
        ("Risks & Open Questions",         memo.get("risks_and_gaps", "")),
        ("Brambles Strategic Fit",         memo.get("brambles_fit", "")),
    ]
    for label, body in sections_data:
        add_label(label)
        add_para(body)
        doc.add_paragraph()

    # Key evidence bullets
    add_label("Key Evidence")
    for item in memo.get("key_evidence", []):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)
    doc.add_paragraph()

    # Recommendation
    add_label("SLAM Recommendation")
    p = doc.add_paragraph()
    run = p.add_run(f"{rec}  —  ")
    run.font.bold = True
    run.font.size = Pt(12)
    rec_colors = {"Pursue": (27, 111, 58), "Monitor": (200, 112, 26), "Pass": (185, 28, 28)}
    r, g, b = rec_colors.get(rec, (27, 58, 75))
    run.font.color.rgb = RGBColor(r, g, b)
    run2 = p.add_run(memo.get("recommendation_paragraph", ""))
    run2.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = re.sub(r"[^\w\-]", "_", row["company_name"])
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_SLAM_Review_Memo.docx"'},
    )
