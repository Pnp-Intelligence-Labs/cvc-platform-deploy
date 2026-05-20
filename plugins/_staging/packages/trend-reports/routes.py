"""
trend-reports plugin routes
============================
Prefix: /reports  (set by plugin_loader from manifest.json)

Endpoints:
  GET    /reports/                          — list all reports
  POST   /reports/                          — create a new report
  GET    /reports/{id}                      — get report + sections + sources
  PATCH  /reports/{id}                      — update title/theme/dates
  DELETE /reports/{id}                      — delete report

  GET    /reports/{id}/sections             — list sections in order
  POST   /reports/{id}/sections             — add a section
  PATCH  /reports/{id}/sections/{sid}       — update section (title/instructions/position/content)
  DELETE /reports/{id}/sections/{sid}       — delete section
  POST   /reports/{id}/sections/{sid}/generate  — trigger section generation (background)
  POST   /reports/{id}/sections/{sid}/reorder   — move section up/down

  GET    /reports/{id}/sources              — list all sources for a report
  POST   /reports/{id}/sources              — add a source (paste/article/db_query)
  POST   /reports/{id}/sources/upload       — upload a PDF source
  PATCH  /reports/{id}/sources/{src_id}     — update label/section assignment
  DELETE /reports/{id}/sources/{src_id}     — delete source

  POST   /reports/{id}/brief                — (re)generate the report brief
  POST   /reports/{id}/publish              — assemble final HTML from all sections
"""

from __future__ import annotations

import json
import logging
import os
import re
import threading
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, UploadFile, File, Form
from pydantic import BaseModel

from api.auth import require_auth
from api.routes.auth import require_jwt, UserInfo
from core.db.connection import get_connection

logger = logging.getLogger(__name__)
router = APIRouter()

# ── Repo root — prefer explicit env var so install path doesn't matter ─────────
_REPO_ROOT = Path(os.environ.get("PLATFORM_ROOT", str(Path(__file__).resolve().parents[3])))

# ── Report upload dir — stored under api/static so the API can serve files ─────
_REPORT_UPLOAD_DIR = str(_REPO_ROOT / "api" / "static" / "report_sources")
os.makedirs(_REPORT_UPLOAD_DIR, exist_ok=True)

# ── Team identity — read from config/team.json at startup ─────────────────────
_TEAM_CONFIG_PATH = _REPO_ROOT / "config" / "team.json"
_team_name = "the venture team"
try:
    with open(_TEAM_CONFIG_PATH) as _f:
        _tc = json.load(_f)
    _team_name = _tc.get("team_name") or _tc.get("team_short") or "the venture team"
except Exception:
    pass


# ── URL content fetcher ──────────────────────────────────────────────────────

def _fetch_url_content(url: str) -> tuple[str, str]:
    """
    Fetch a URL and extract readable text. Returns (content_text, status_note).
    Handles arxiv abstract/HTML pages specially. Falls back to generic HTML strip.
    Returns ("", error_note) on failure — never raises.
    """
    import re
    import requests as _req

    headers = {
        'User-Agent': 'Mozilla/5.0 (compatible; PlatformBot/1.0)',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
    }

    try:
        resp = _req.get(url, timeout=12, headers=headers, allow_redirects=True)
        resp.raise_for_status()
    except Exception as e:
        return "", f"Fetch failed: {e}"

    content_type = resp.headers.get('content-type', '')
    if 'pdf' in content_type:
        return "", "PDF at URL — upload as file instead of URL for full text extraction"

    html = resp.text

    # Try BeautifulSoup (preferred)
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        # Remove noise elements
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside',
                         'noscript', 'form', 'button', 'iframe']):
            tag.decompose()

        # arxiv: prioritise abstract + title
        if 'arxiv.org/abs/' in url:
            title_el = soup.find('h1', class_='title')
            abstract_el = soup.find('blockquote', class_='abstract')
            title_txt = title_el.get_text(strip=True).replace('Title:', '').strip() if title_el else ''
            abstract_txt = abstract_el.get_text(strip=True).replace('Abstract:', '').strip() if abstract_el else ''
            if title_txt or abstract_txt:
                return f"Title: {title_txt}\n\nAbstract: {abstract_txt}", "arxiv abstract"

        # arxiv HTML full paper
        if 'arxiv.org/html/' in url:
            body = soup.find('article') or soup.find('main') or soup.body
            text = body.get_text(separator='\n', strip=True) if body else soup.get_text('\n', strip=True)
            text = re.sub(r'\n{3,}', '\n\n', text)
            return text[:30000], f"arxiv HTML — {len(text)} chars"

        # Nature / ScienceDirect / Frontiers — grab article body
        article = soup.find('article') or soup.find('main') or soup.find(id='main-content') or soup.body
        text = article.get_text(separator='\n', strip=True) if article else soup.get_text('\n', strip=True)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text[:30000], f"{len(text)} chars extracted"

    except ImportError:
        # bs4 not available — regex strip
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text).strip()
        return text[:30000], f"regex strip — {len(text)} chars"


# ── Pydantic models ──────────────────────────────────────────────────────────

class CreateReportRequest(BaseModel):
    title: str
    sector: Optional[str] = None
    theme: Optional[str] = None
    date_from: Optional[str] = None
    date_to: Optional[str] = None
    output_format: Optional[str] = 'report'   # 'report' | 'blog'
    citation_style: Optional[str] = 'superscript'
    audience: Optional[str] = 'practitioner'   # executive | practitioner | investor | analyst | general
    tone: Optional[str] = 'analytical'          # analytical | authoritative | narrative | concise | conversational


class PatchReportRequest(BaseModel):
    title: Optional[str] = None
    sector: Optional[str] = None
    theme: Optional[str] = None
    date_from: Optional[str] = None
    date_to: Optional[str] = None
    status: Optional[str] = None
    report_brief: Optional[str] = None
    output_format: Optional[str] = None
    citation_style: Optional[str] = None
    published_html: Optional[str] = None
    audience: Optional[str] = None
    tone: Optional[str] = None


class CreateSectionRequest(BaseModel):
    title: str
    instructions: Optional[str] = None
    data_sources: list = []
    section_type: Optional[str] = 'prose'
    audience: Optional[str] = None   # None = inherit from report
    tone: Optional[str] = None       # None = inherit from report


class PatchSectionRequest(BaseModel):
    title: Optional[str] = None
    instructions: Optional[str] = None
    data_sources: Optional[list] = None
    position: Optional[int] = None
    content: Optional[str] = None    # manual edit of generated content
    section_type: Optional[str] = None
    audience: Optional[str] = None   # '' sentinel = clear override (use report default)
    tone: Optional[str] = None       # '' sentinel = clear override


class CreateSourceRequest(BaseModel):
    source_type: str           # paste | article | db_query
    label: Optional[str] = None
    section_id: Optional[int] = None
    content_text: Optional[str] = None
    query_sql: Optional[str] = None
    article_url: Optional[str] = None
    chart_type: Optional[str] = None   # bar | line | pie | area — for chart injection in assembled HTML
    x_key: Optional[str] = None
    y_key: Optional[str] = None


class PatchSourceRequest(BaseModel):
    label: Optional[str] = None
    section_id: Optional[int] = None   # None = shared across all sections


# ── Helpers ──────────────────────────────────────────────────────────────────

def _get_report_or_404(report_id: int):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM cvc.trend_reports WHERE id = %s",
                (report_id,)
            )
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Report not found")
    return dict(row)


def _get_section_or_404(report_id: int, section_id: int):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM cvc.report_sections WHERE id = %s AND report_id = %s",
                (section_id, report_id)
            )
            row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Section not found")
    return dict(row)


def _serialize(row: dict) -> dict:
    """Convert non-JSON-serializable types."""
    import datetime
    from decimal import Decimal
    import uuid
    out = {}
    for k, v in row.items():
        if isinstance(v, (datetime.datetime, datetime.date)):
            out[k] = v.isoformat()
        elif isinstance(v, Decimal):
            out[k] = float(v)
        elif isinstance(v, uuid.UUID):
            out[k] = str(v)
        else:
            out[k] = v
    return out


# ── Report CRUD ──────────────────────────────────────────────────────────────

@router.get("/")
def list_reports(user: UserInfo = Depends(require_jwt)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT r.*,
                       COUNT(s.id) AS section_count
                FROM cvc.trend_reports r
                LEFT JOIN cvc.report_sections s ON s.report_id = r.id
                GROUP BY r.id
                ORDER BY r.created_at DESC
            """)
            rows = [_serialize(dict(row)) for row in cur.fetchall()]
    return rows


@router.post("/")
def create_report(req: CreateReportRequest, user: UserInfo = Depends(require_jwt)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO cvc.trend_reports (title, sector, theme, date_from, date_to, created_by, output_format, citation_style, audience, tone)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING *
            """, (req.title, req.sector, req.theme, req.date_from, req.date_to, user.username,
                  req.output_format or 'report', req.citation_style or 'superscript',
                  req.audience or 'practitioner', req.tone or 'analytical'))
            row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.get("/catalog")
def get_catalog(user: UserInfo = Depends(require_jwt)):
    """Return a hardcoded catalog of pre-built data suggestions grouped by category."""
    return [
        {
            "name": "Funding Activity",
            "suggestions": [
                {
                    "title": "Funding Rounds by Sector (Last 2 Years)",
                    "rationale": "Shows where capital has been flowing — useful for market momentum sections",
                    "sql": "SELECT c.sector, COUNT(fr.id) AS round_count, SUM(fr.amount_usd) AS total_usd FROM cvc.funding_rounds fr JOIN cvc.companies c ON c.id = fr.company_id WHERE fr.announced_date >= NOW() - INTERVAL '2 years' GROUP BY c.sector ORDER BY total_usd DESC NULLS LAST",
                    "chart_type": "bar",
                    "x_key": "sector",
                    "y_key": "total_usd",
                    "description": "Total funding raised per sector over the last 24 months",
                },
                {
                    "title": "Round Type Distribution",
                    "rationale": "Shows whether the market is early-stage heavy or scaling — good for investment thesis framing",
                    "sql": "SELECT round_type, COUNT(*) AS count FROM cvc.funding_rounds WHERE announced_date >= NOW() - INTERVAL '2 years' GROUP BY round_type ORDER BY count DESC",
                    "chart_type": "pie",
                    "x_key": "round_type",
                    "y_key": "count",
                    "description": "Distribution of funding rounds by type over the last 24 months",
                },
                {
                    "title": "Funding Activity by Quarter",
                    "rationale": "Trend line showing acceleration or slowdown — useful for macro framing",
                    "sql": "SELECT TO_CHAR(DATE_TRUNC('quarter', announced_date), 'YYYY-Q') AS quarter, COUNT(*) AS rounds, SUM(amount_usd) AS total_usd FROM cvc.funding_rounds WHERE announced_date IS NOT NULL GROUP BY quarter ORDER BY quarter",
                    "chart_type": "line",
                    "x_key": "quarter",
                    "y_key": "total_usd",
                    "description": "Quarterly funding volume trend across all sectors",
                },
                {
                    "title": "Top 20 Funded Companies",
                    "rationale": "Identifies the best-capitalized players in your coverage universe",
                    "sql": "SELECT name, sector, funding_total_usd, stage FROM cvc.companies WHERE funding_total_usd IS NOT NULL ORDER BY funding_total_usd DESC LIMIT 20",
                    "chart_type": "bar",
                    "x_key": "name",
                    "y_key": "funding_total_usd",
                    "description": "Companies ranked by total funding raised",
                },
            ],
        },
        {
            "name": "Portfolio & Pipeline",
            "suggestions": [
                {
                    "title": "Portfolio Companies by Sector",
                    "rationale": "Shows portfolio concentration — good for portfolio construction sections",
                    "sql": "SELECT sector, COUNT(*) AS count FROM cvc.companies WHERE is_portfolio = TRUE GROUP BY sector ORDER BY count DESC",
                    "chart_type": "bar",
                    "x_key": "sector",
                    "y_key": "count",
                    "description": "Portfolio company distribution across sectors",
                },
                {
                    "title": "Pipeline by Stage",
                    "rationale": "Shows funnel shape from sourcing to investment — useful for sourcing coverage analysis",
                    "sql": "SELECT stage, COUNT(*) AS count FROM cvc.companies WHERE is_portfolio = FALSE GROUP BY stage ORDER BY count DESC",
                    "chart_type": "bar",
                    "x_key": "stage",
                    "y_key": "count",
                    "description": "Non-portfolio companies in the pipeline by investment stage",
                },
                {
                    "title": "Composite Score Distribution",
                    "rationale": "Shows the quality spread of the pipeline — useful for investment readiness sections",
                    "sql": "SELECT CASE WHEN score_composite >= 8 THEN '8-10 (High)' WHEN score_composite >= 6 THEN '6-8 (Above Avg)' WHEN score_composite >= 4 THEN '4-6 (Average)' ELSE '0-4 (Low)' END AS score_band, COUNT(*) AS count FROM cvc.companies WHERE score_composite IS NOT NULL GROUP BY score_band ORDER BY score_band",
                    "chart_type": "bar",
                    "x_key": "score_band",
                    "y_key": "count",
                    "description": "Distribution of composite readiness scores across all tracked companies",
                },
                {
                    "title": "Top Scoring Companies by Sector",
                    "rationale": "Identifies the highest-quality companies per vertical for competitive analysis",
                    "sql": "SELECT name, sector, score_composite, stage FROM cvc.companies WHERE score_composite IS NOT NULL ORDER BY score_composite DESC LIMIT 30",
                    "chart_type": "bar",
                    "x_key": "name",
                    "y_key": "score_composite",
                    "description": "Top 30 companies ranked by composite readiness score",
                },
            ],
        },
        {
            "name": "Sector Signals",
            "suggestions": [
                {
                    "title": "Signal Volume by Sector",
                    "rationale": "Shows which sectors are generating the most market noise — useful for trend identification",
                    "sql": "SELECT sector, COUNT(*) AS signal_count FROM cvc.raw_signals WHERE sector IS NOT NULL GROUP BY sector ORDER BY signal_count DESC",
                    "chart_type": "bar",
                    "x_key": "sector",
                    "y_key": "signal_count",
                    "description": "Total market signals collected per sector",
                },
                {
                    "title": "Signal Types Distribution",
                    "rationale": "Shows the nature of market activity — partnership vs product vs regulatory",
                    "sql": "SELECT signal_type, COUNT(*) AS count FROM cvc.raw_signals WHERE signal_type IS NOT NULL GROUP BY signal_type ORDER BY count DESC LIMIT 15",
                    "chart_type": "pie",
                    "x_key": "signal_type",
                    "y_key": "count",
                    "description": "Distribution of signal types (partnerships, product launches, regulatory, etc.)",
                },
                {
                    "title": "Signal Activity by Quarter",
                    "rationale": "Trend line for market intelligence activity over time",
                    "sql": "SELECT quarter, COUNT(*) AS count FROM cvc.raw_signals WHERE quarter IS NOT NULL GROUP BY quarter ORDER BY quarter",
                    "chart_type": "line",
                    "x_key": "quarter",
                    "y_key": "count",
                    "description": "Weekly market signals collected per quarter",
                },
                {
                    "title": "High-Confidence Briefing Insights by Sector",
                    "rationale": "Surfaces the most reliable intelligence signals for a thesis-grounding section",
                    "sql": "SELECT sector, COUNT(*) AS count FROM cvc.briefing_insights WHERE confidence = 'HIGH' AND sector IS NOT NULL GROUP BY sector ORDER BY count DESC",
                    "chart_type": "bar",
                    "x_key": "sector",
                    "y_key": "count",
                    "description": "High-confidence enriched insights from podcasts and research articles, by sector",
                },
            ],
        },
        {
            "name": "Market Coverage",
            "suggestions": [
                {
                    "title": "Companies by Founded Year",
                    "rationale": "Shows the vintage distribution of the startup ecosystem — is it mature or early?",
                    "sql": "SELECT founded, COUNT(*) AS count FROM cvc.companies WHERE founded IS NOT NULL AND founded >= 2010 GROUP BY founded ORDER BY founded",
                    "chart_type": "bar",
                    "x_key": "founded",
                    "y_key": "count",
                    "description": "Number of companies founded per year since 2010",
                },
                {
                    "title": "Geographic Distribution (Top 15 Countries)",
                    "rationale": "Shows where the innovation is concentrated — useful for global market framing",
                    "sql": "SELECT hq_country, COUNT(*) AS count FROM cvc.companies WHERE hq_country IS NOT NULL GROUP BY hq_country ORDER BY count DESC LIMIT 15",
                    "chart_type": "bar",
                    "x_key": "hq_country",
                    "y_key": "count",
                    "description": "Top 15 countries by number of tracked companies",
                },
                {
                    "title": "Partner Introductions by Outcome",
                    "rationale": "Shows commercial traction of partner network — useful for partner integration sections",
                    "sql": "SELECT outcome, COUNT(*) AS count FROM cvc.partner_intros WHERE outcome IS NOT NULL GROUP BY outcome ORDER BY count DESC",
                    "chart_type": "pie",
                    "x_key": "outcome",
                    "y_key": "count",
                    "description": "Outcomes of partner introductions across the CVC ecosystem",
                },
                {
                    "title": "Employee Count Distribution",
                    "rationale": "Shows team size maturity — are these pre-revenue startups or scaling businesses?",
                    "sql": "SELECT CASE WHEN employee_count < 10 THEN '<10' WHEN employee_count < 50 THEN '10-49' WHEN employee_count < 200 THEN '50-199' WHEN employee_count < 500 THEN '200-499' ELSE '500+' END AS size_band, COUNT(*) AS count FROM cvc.companies WHERE employee_count IS NOT NULL GROUP BY size_band ORDER BY size_band",
                    "chart_type": "bar",
                    "x_key": "size_band",
                    "y_key": "count",
                    "description": "Companies grouped by employee headcount range",
                },
            ],
        },
    ]


@router.get("/{report_id}")
def get_report(report_id: int, user: UserInfo = Depends(require_jwt)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM cvc.trend_reports WHERE id = %s", (report_id,))
            r = cur.fetchone()
            if not r:
                raise HTTPException(status_code=404, detail="Report not found")
            report = _serialize(dict(r))

            cur.execute("""
                SELECT * FROM cvc.report_sections
                WHERE report_id = %s ORDER BY position, id
            """, (report_id,))
            sections = [_serialize(dict(row)) for row in cur.fetchall()]

            cur.execute("""
                SELECT * FROM cvc.report_sources
                WHERE report_id = %s ORDER BY created_at
            """, (report_id,))
            sources = [_serialize(dict(row)) for row in cur.fetchall()]

    report['sections'] = sections
    report['sources'] = sources
    return report


@router.patch("/{report_id}")
def patch_report(report_id: int, req: PatchReportRequest, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    if not updates:
        return {"ok": True}
    set_clause = ", ".join(f"{k} = %s" for k in updates)
    values = list(updates.values()) + [report_id]
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE cvc.trend_reports SET {set_clause}, updated_at = NOW() WHERE id = %s RETURNING *",
                values
            )
            row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.delete("/{report_id}")
def delete_report(report_id: int, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM cvc.trend_reports WHERE id = %s", (report_id,))
        conn.commit()
    return {"ok": True}


# ── Section CRUD ─────────────────────────────────────────────────────────────

@router.get("/{report_id}/sections")
def list_sections(report_id: int, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT * FROM cvc.report_sections
                WHERE report_id = %s ORDER BY position, id
            """, (report_id,))
            return [_serialize(dict(row)) for row in cur.fetchall()]


@router.post("/{report_id}/sections")
def create_section(report_id: int, req: CreateSectionRequest, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT COALESCE(MAX(position), -1) + 1 AS next_pos FROM cvc.report_sections WHERE report_id = %s
            """, (report_id,))
            next_pos = cur.fetchone()['next_pos']
            cur.execute("""
                INSERT INTO cvc.report_sections (report_id, position, title, instructions, data_sources, section_type, audience, tone)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING *
            """, (report_id, next_pos, req.title, req.instructions, json.dumps(req.data_sources),
                  req.section_type or 'prose', req.audience or None, req.tone or None))
            row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.patch("/{report_id}/sections/{section_id}")
def patch_section(report_id: int, section_id: int, req: PatchSectionRequest,
                  user: UserInfo = Depends(require_jwt)):
    _get_section_or_404(report_id, section_id)
    updates: dict = {}
    if req.title is not None:
        updates['title'] = req.title
    if req.instructions is not None:
        updates['instructions'] = req.instructions
    if req.data_sources is not None:
        updates['data_sources'] = json.dumps(req.data_sources)
    if req.position is not None:
        updates['position'] = req.position
    if req.section_type is not None:
        updates['section_type'] = req.section_type
    # audience/tone: '' = clear override (null), non-empty = set, not in fields_set = skip
    if 'audience' in req.model_fields_set:
        updates['audience'] = req.audience if req.audience else None
    if 'tone' in req.model_fields_set:
        updates['tone'] = req.tone if req.tone else None
    if req.content is not None:
        updates['content'] = req.content
        # Push to version history on manual edit
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT version_history, generated_at, confidence_score FROM cvc.report_sections WHERE id = %s", (section_id,))
                existing = dict(cur.fetchone())
            conn.commit()
        history = existing.get('version_history') or []
        if existing.get('content'):
            import datetime
            history.append({
                'content': existing.get('content'),
                'generated_at': existing.get('generated_at').isoformat() if existing.get('generated_at') else None,
                'confidence_score': existing.get('confidence_score'),
            })
        updates['version_history'] = json.dumps(history)

    if not updates:
        return {"ok": True}
    set_clause = ", ".join(f"{k} = %s" for k in updates)
    values = list(updates.values()) + [section_id]
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE cvc.report_sections SET {set_clause} WHERE id = %s RETURNING *",
                values
            )
            row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.delete("/{report_id}/sections/{section_id}")
def delete_section(report_id: int, section_id: int, user: UserInfo = Depends(require_jwt)):
    _get_section_or_404(report_id, section_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM cvc.report_sections WHERE id = %s", (section_id,))
        conn.commit()
    return {"ok": True}


# ── Section Generation ────────────────────────────────────────────────────────

@router.post("/{report_id}/generate-outline")
def generate_outline(report_id: int, user: UserInfo = Depends(require_jwt)):
    """Use LLM to generate 5-7 suggested sections for the report. Returns draft — does NOT save."""
    import requests as _req

    report = _get_report_or_404(report_id)
    _key = os.environ.get("OPENROUTER_REPORT_KEY") or os.environ.get("OPENROUTER_API_KEY", "")

    _VOICE_PREFIX = (
        f"You are writing venture intelligence reports for {_team_name}. "
        "Write in a confident, direct, practitioner voice. "
        "Lead every paragraph with the claim. Use active voice."
    )

    system = (
        _VOICE_PREFIX + "\n\n"
        "You are a research analyst planning a venture intelligence report. "
        "Given the report context, return ONLY a JSON array of section objects. "
        "No markdown fences, no explanation — just the JSON array.\n\n"
        "Each element must have exactly:\n"
        '{"title": "<short section title, 5 words max>", "instructions": "<1-2 sentences of specific analytical guidance for this section>", "section_type": "<type>"}\n\n'
        "Section type guide — choose the best fit for each section:\n"
        "  prose          — standard body section (default for most sections)\n"
        "  deep_dive      — extended technical analysis, 5-7 paragraphs\n"
        "  sidebar        — a 2-3 sentence highlighted factoid or key stat\n"
        "  spotlight      — a 150-word featured company or technology profile\n"
        "  tech_stack     — vendor landscape and integration patterns\n"
        "  investment_take — market sizing, valuation, portfolio fit\n\n"
        "A good report mixes types. Include at least one investment_take and one tech_stack if the topic warrants it. "
        "Return 5-7 sections that together form a complete, coherent report structure."
    )

    user_msg = (
        f"Report title: {report.get('title', '')}\n"
        f"Sector: {report.get('sector') or 'Not specified'}\n"
        f"Theme: {report.get('theme') or 'Not specified'}\n"
        f"Brief: {report.get('report_brief') or 'Not written yet'}\n\n"
        "Generate 5-7 section titles and instructions for this report."
    )

    try:
        resp = _req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {_key}", "Content-Type": "application/json"},
            json={
                "model": "moonshotai/kimi-k2.6",
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user_msg},
                ],
                "temperature": 0.3,
            },
            timeout=60,
        )
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()
        # Strip <think> blocks (Qwen3 reasoning models)
        import re as _re
        raw = _re.sub(r'<think>[\s\S]*?</think>', '', raw, flags=_re.IGNORECASE).strip()
        # Strip markdown code fences
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        sections = json.loads(raw.strip())
        if not isinstance(sections, list):
            raise ValueError("LLM did not return a JSON array")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM outline generation failed: {e}")

    return {"sections": sections}


class BulkCreateSectionsRequest(BaseModel):
    sections: list


@router.post("/{report_id}/sections/bulk-create")
def bulk_create_sections(report_id: int, req: BulkCreateSectionsRequest,
                         user: UserInfo = Depends(require_jwt)):
    """Save a list of {title, instructions} sections in order."""
    _get_report_or_404(report_id)
    created = []
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT COALESCE(MAX(position), -1) AS max_pos FROM cvc.report_sections WHERE report_id = %s",
                (report_id,)
            )
            start_pos = cur.fetchone()['max_pos'] + 1
            for i, s in enumerate(req.sections):
                cur.execute("""
                    INSERT INTO cvc.report_sections (report_id, position, title, instructions, data_sources, section_type)
                    VALUES (%s, %s, %s, %s, '[]', %s)
                    RETURNING *
                """, (report_id, start_pos + i, s.get('title', f'Section {i+1}'), s.get('instructions'), s.get('section_type', 'prose')))
                created.append(_serialize(dict(cur.fetchone())))
        conn.commit()
    return {"sections": created}


@router.post("/{report_id}/sections/{section_id}/discover")
def discover_sources(report_id: int, section_id: int, user: UserInfo = Depends(require_jwt)):
    """
    Generate search queries from the section title + report brief,
    run them against Brave Search, and query the DB for matching companies + insights.
    Returns candidates for the analyst to select from.
    """
    import requests as _req
    import re as _re

    report = _get_report_or_404(report_id)
    section = _get_section_or_404(report_id, section_id)

    _key = os.environ.get("OPENROUTER_REPORT_KEY") or os.environ.get("OPENROUTER_API_KEY", "")
    BRAVE_KEY = os.environ.get("BRAVE_SEARCH_KEY", "")

    sector = report.get("sector") or ""
    section_title = section.get("title", "")
    report_brief = report.get("report_brief") or ""
    report_title = report.get("title", "")

    # ── Step 1: Generate search queries ──────────────────────────────────────
    query_prompt = (
        f"Report: {report_title}\nSector: {sector}\nBrief: {report_brief}\n"
        f"Section: {section_title}\nInstructions: {section.get('instructions') or ''}\n\n"
        "Generate 2-3 focused web search queries that would find the best external research "
        "for this section. Return ONLY a JSON array of query strings. No explanation."
    )
    search_queries: list[str] = []
    try:
        resp = _req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {_key}", "Content-Type": "application/json"},
            json={
                "model": "moonshotai/kimi-k2.6",
                "messages": [{"role": "user", "content": query_prompt}],
                "temperature": 0.2,
            },
            timeout=30,
        )
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()
        raw = _re.sub(r'<think>[\s\S]*?</think>', '', raw, flags=_re.IGNORECASE).strip()
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        search_queries = json.loads(raw.strip())
        if not isinstance(search_queries, list):
            search_queries = []
    except Exception as e:
        logger.warning(f"Query generation failed for section {section_id}: {e}")
        # Fallback: use section title + sector as query
        search_queries = [f"{section_title} {sector} research 2024", f"{sector} {report_title}"]

    # ── Step 2: Run Brave searches ────────────────────────────────────────────
    web_sources: list[dict] = []
    seen_urls: set[str] = set()

    if BRAVE_KEY:
        for query in search_queries[:3]:
            try:
                brave_resp = _req.get(
                    "https://api.search.brave.com/res/v1/web/search",
                    params={"q": query, "count": 5, "text_decorations": False},
                    headers={"Accept": "application/json", "X-Subscription-Token": BRAVE_KEY},
                    timeout=10,
                )
                results = brave_resp.json().get("web", {}).get("results", [])
                for r in results:
                    url = r.get("url", "")
                    if url and url not in seen_urls:
                        seen_urls.add(url)
                        web_sources.append({
                            "title": r.get("title", url),
                            "url": url,
                            "snippet": r.get("description", ""),
                        })
            except Exception as e:
                logger.warning(f"Brave search failed for query '{query}': {e}")
    else:
        logger.info("BRAVE_SEARCH_KEY not set — skipping web search for discover")

    # ── Step 3: DB sources ────────────────────────────────────────────────────
    db_sources: list[dict] = []
    companies_sql = None
    insights_sql = None

    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                if sector:
                    companies_sql = (
                        "SELECT name, stage, score_composite, one_liner FROM cvc.companies "
                        f"WHERE sector = '{sector}' AND score_composite IS NOT NULL "
                        "ORDER BY score_composite DESC LIMIT 20"
                    )
                    cur.execute(
                        "SELECT name, stage, score_composite, one_liner FROM cvc.companies "
                        "WHERE sector = %s AND score_composite IS NOT NULL "
                        "ORDER BY score_composite DESC LIMIT 20",
                        (sector,)
                    )
                    companies = [dict(r) for r in cur.fetchall()]
                    if companies:
                        preview = [c['name'] for c in companies[:5]]
                        db_sources.append({
                            "id": "companies",
                            "label": f"Top {sector} companies by score",
                            "description": f"{len(companies)} companies",
                            "sql": companies_sql,
                            "preview": preview,
                        })

                    insights_sql = (
                        "SELECT insight, expert, confidence, source_title FROM cvc.briefing_insights "
                        f"WHERE sector = '{sector}' ORDER BY created_at DESC LIMIT 10"
                    )
                    cur.execute(
                        "SELECT insight, expert, confidence, source_title FROM cvc.briefing_insights "
                        "WHERE sector = %s ORDER BY created_at DESC LIMIT 10",
                        (sector,)
                    )
                    insights = [dict(r) for r in cur.fetchall()]
                    if insights:
                        preview = [i['insight'][:80] + ('...' if len(i['insight']) > 80 else '') for i in insights[:3]]
                        db_sources.append({
                            "id": "insights",
                            "label": f"Recent briefing insights — {sector}",
                            "description": f"{len(insights)} items",
                            "sql": insights_sql,
                            "preview": preview,
                        })
                else:
                    # No sector — use generic company query
                    companies_sql = (
                        "SELECT name, sector, stage, score_composite FROM cvc.companies "
                        "WHERE score_composite IS NOT NULL ORDER BY score_composite DESC LIMIT 20"
                    )
                    cur.execute(
                        "SELECT name, sector, stage, score_composite FROM cvc.companies "
                        "WHERE score_composite IS NOT NULL ORDER BY score_composite DESC LIMIT 20"
                    )
                    companies = [dict(r) for r in cur.fetchall()]
                    if companies:
                        preview = [c['name'] for c in companies[:5]]
                        db_sources.append({
                            "id": "companies",
                            "label": "Top companies by score (all sectors)",
                            "description": f"{len(companies)} companies",
                            "sql": companies_sql,
                            "preview": preview,
                        })
    except Exception as e:
        logger.error(f"DB source discovery failed for section {section_id}: {e}")

    return {
        "web_sources": web_sources,
        "db_sources": db_sources,
        "queries_used": search_queries,
        "brave_available": bool(BRAVE_KEY),
    }


@router.post("/{report_id}/sections/{section_id}/generate")
def generate_section(report_id: int, section_id: int, background_tasks: BackgroundTasks,
                     user: UserInfo = Depends(require_jwt)):
    """Trigger async LLM generation for a single section."""
    report = _get_report_or_404(report_id)
    section = _get_section_or_404(report_id, section_id)

    # Mark as generating
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.report_sections SET status = 'generating', error_msg = NULL WHERE id = %s",
                (section_id,)
            )
        conn.commit()

    background_tasks.add_task(_generate_section_bg, report_id, section_id)
    return {"ok": True, "status": "generating"}


def _generate_section_bg(report_id: int, section_id: int):
    """Background task: assemble sources, call LLM, apply voice pass, save."""
    try:
        import sys
        sys.path.insert(0, str(_REPO_ROOT))
        from workers.trends.report_builder import build_section
        build_section(report_id, section_id)
    except Exception as e:
        logger.error(f"Section generation failed [{section_id}]: {e}", exc_info=True)
        try:
            with get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE cvc.report_sections SET status = 'error', error_msg = %s WHERE id = %s",
                        (str(e)[:500], section_id)
                    )
                conn.commit()
        except Exception:
            pass


# ── Sources ───────────────────────────────────────────────────────────────────

@router.get("/{report_id}/sources")
def list_sources(report_id: int, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM cvc.report_sources WHERE report_id = %s ORDER BY created_at",
                (report_id,)
            )
            return [_serialize(dict(row)) for row in cur.fetchall()]


@router.post("/{report_id}/sources")
def create_source(report_id: int, req: CreateSourceRequest, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)

    content_text = req.content_text
    fetch_status = None
    query_result = None

    # ── URL article: fetch and extract content ────────────────────────────────
    if req.source_type == 'article':
        url = req.article_url or (req.label if req.label and req.label.startswith('http') else None)
        if url and not content_text:
            fetched, fetch_status = _fetch_url_content(url)
            if fetched:
                content_text = fetched
                logger.info(f"URL fetch OK [{report_id}]: {url} — {fetch_status}")
            else:
                logger.warning(f"URL fetch failed [{report_id}]: {url} — {fetch_status}")

        # Save article to content_items so the briefing pipeline can enrich it
        if url:
            try:
                import hashlib
                from urllib.parse import urlparse as _urlparse
                content_hash = hashlib.sha256(url.encode()).hexdigest()
                source_domain = _urlparse(url).netloc
                item_title = (req.label or url)[:1000]
                with get_connection() as _conn:
                    with _conn.cursor() as _cur:
                        _cur.execute("""
                            INSERT INTO cvc.content_items
                                (content_type, title, url, source, raw_text, enrichment_status, content_hash)
                            VALUES ('article', %s, %s, %s, %s, 'raw', %s)
                            ON CONFLICT (content_hash) DO NOTHING
                        """, (item_title, url[:2000], source_domain[:500],
                              (content_text or '')[:30000], content_hash))
                    _conn.commit()
                logger.info(f"Article saved to content_items: {url}")
            except Exception as _e:
                logger.warning(f"content_items insert skipped for {url}: {_e}")

    # ── DB query: run and store result ────────────────────────────────────────
    if req.source_type == 'db_query' and req.query_sql:
        try:
            with get_connection() as conn:
                with conn.cursor() as cur:
                    sql = req.query_sql.strip()
                    if any(kw in sql.upper() for kw in ('INSERT', 'UPDATE', 'DELETE', 'DROP', 'TRUNCATE', 'ALTER')):
                        raise HTTPException(status_code=400, detail="Only SELECT queries are allowed")
                    cur.execute(sql)
                    rows = cur.fetchmany(200)
                    # RealDictCursor returns dict rows — use dict(row) not zip
                    query_result = json.dumps([dict(row) for row in rows], default=str)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Query error: {e}")

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO cvc.report_sources
                    (report_id, section_id, source_type, label, content_text, query_sql, query_result, article_url,
                     chart_type, x_key, y_key)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING *
            """, (
                report_id, req.section_id, req.source_type, req.label,
                content_text, req.query_sql,
                query_result,
                req.article_url or (req.label if req.label and req.label.startswith('http') else None),
                req.chart_type, req.x_key, req.y_key,
            ))
            row = _serialize(dict(cur.fetchone()))
        conn.commit()

    # Surface fetch status to caller so UI can show "fetched N chars" or warn
    if fetch_status:
        row['fetch_status'] = fetch_status
    return row


@router.post("/{report_id}/sources/upload")
async def upload_source(
    report_id: int,
    file: UploadFile = File(...),
    label: str = Form(default=""),
    section_id: Optional[int] = Form(default=None),
    user: UserInfo = Depends(require_jwt),
):
    """Upload a PDF and extract its text."""
    _get_report_or_404(report_id)

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files supported")

    # Save file
    safe_name = f"{report_id}_{file.filename.replace(' ', '_')}"
    file_path = os.path.join(_REPORT_UPLOAD_DIR, safe_name)
    content = await file.read()
    with open(file_path, 'wb') as f:
        f.write(content)

    # Extract text
    content_text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for p in pdf.pages[:60]:   # cap at 60 pages
                t = p.extract_text()
                if t:
                    pages.append(t)
            content_text = "\n\n".join(pages)
    except Exception as e:
        logger.warning(f"PDF extraction failed for {safe_name}: {e}")
        content_text = f"[PDF extraction failed: {e}]"

    disp_label = label or file.filename

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO cvc.report_sources
                    (report_id, section_id, source_type, label, filename, file_path, content_text)
                VALUES (%s, %s, 'pdf', %s, %s, %s, %s)
                RETURNING *
            """, (report_id, section_id or None, disp_label, file.filename, file_path, content_text))
            row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.patch("/{report_id}/sources/{source_id}")
def patch_source(report_id: int, source_id: int, req: PatchSourceRequest,
                 user: UserInfo = Depends(require_jwt)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM cvc.report_sources WHERE id = %s AND report_id = %s",
                        (source_id, report_id))
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Source not found")
            updates: dict = {}
            if req.label is not None:
                updates['label'] = req.label
            if 'section_id' in req.model_fields_set:
                updates['section_id'] = req.section_id
            if updates:
                set_clause = ", ".join(f"{k} = %s" for k in updates)
                cur.execute(
                    f"UPDATE cvc.report_sources SET {set_clause} WHERE id = %s RETURNING *",
                    list(updates.values()) + [source_id]
                )
                row = _serialize(dict(cur.fetchone()))
            else:
                cur.execute("SELECT * FROM cvc.report_sources WHERE id = %s", (source_id,))
                row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.delete("/{report_id}/sources/{source_id}")
def delete_source(report_id: int, source_id: int, user: UserInfo = Depends(require_jwt)):
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT file_path FROM cvc.report_sources WHERE id = %s AND report_id = %s",
                        (source_id, report_id))
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Source not found")
            if row['file_path']:
                try:
                    os.remove(row['file_path'])
                except FileNotFoundError:
                    pass
            cur.execute("DELETE FROM cvc.report_sources WHERE id = %s", (source_id,))
        conn.commit()
    return {"ok": True}


# ── Import Outline ────────────────────────────────────────────────────────────

class ImportOutlineRequest(BaseModel):
    content_text: str   # raw text — either pasted outline or pre-extracted PDF text


@router.post("/{report_id}/import-outline")
async def import_outline(report_id: int, req: ImportOutlineRequest,
                         user: UserInfo = Depends(require_jwt)):
    """
    Given raw text (pasted outline or extracted PDF), use the LLM to extract:
      - A list of {title, instructions} section objects
      - An optional report brief (only set if the report doesn't have one yet)
    Creates all sections in order and returns the updated report.
    """
    report = _get_report_or_404(report_id)

    import requests as _req
    _key = os.environ.get("OPENROUTER_REPORT_KEY") or os.environ.get("OPENROUTER_API_KEY", "")

    system = """You are a report structure analyst. Given raw text — which may be a document, an outline, a table of contents, or rough notes — extract a clean report outline.

Return ONLY valid JSON matching this exact schema:
{
  "brief": "<2-3 sentence brief describing the report's argument, audience, and angle — or null if the text doesn't provide enough context>",
  "sections": [
    {"title": "<section title>", "instructions": "<instructions>", "section_type": "<type>"},
    ...
  ]
}

Section type rules — read the input carefully for explicit type labels (e.g. "Type: Sidebar", "Type: Deep Dive") and use them exactly. If no type is specified, infer the best fit:
  prose          — standard body section (default)
  deep_dive      — extended technical analysis
  sidebar        — short 2-3 sentence callout or factoid box
  spotlight      — featured company or technology profile
  tech_stack     — vendor landscape and integration patterns
  investment_take — market sizing, valuation, portfolio fit

Instructions rules:
- Preserve ALL sub-topics, bullet points, and specific details from the input — fold them into the instructions as a rich single paragraph.
- Do not compress or genericize. If the input says "cover Harmonic Drive gearing and GelSight tactile sensing", those exact terms must appear in the instructions.
- Capture the specific angle or argument, not just a restatement of the title.

General rules:
- Extract between 3 and 12 sections. If the document has more, group minor ones.
- Section titles should be short and declarative (5 words max).
- If the input is a finished document, derive the outline from what each section actually covers.
- Return nothing outside the JSON block."""

    user_msg = f"""Report title: {report.get('title', '')}
Sector: {report.get('sector') or 'Not specified'}

Content to extract outline from:
---
{req.content_text[:12000]}
---

Extract the outline now."""

    try:
        resp = _req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {_key}", "Content-Type": "application/json"},
            json={
                "model": "moonshotai/kimi-k2.6",
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user_msg},
                ],
                "temperature": 0.1,
            },
            timeout=60,
        )
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()
        # Strip markdown code fences if present
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        parsed = json.loads(raw.strip())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM extraction failed: {e}")

    sections_data = parsed.get("sections") or []
    brief_text    = parsed.get("brief")

    if not sections_data:
        raise HTTPException(status_code=422, detail="Could not extract any sections from the provided content")

    created_sections = []
    with get_connection() as conn:
        with conn.cursor() as cur:
            # Get current max position
            cur.execute(
                "SELECT COALESCE(MAX(position), -1) AS max_pos FROM cvc.report_sections WHERE report_id = %s",
                (report_id,)
            )
            start_pos = cur.fetchone()['max_pos'] + 1

            for i, s in enumerate(sections_data):
                cur.execute("""
                    INSERT INTO cvc.report_sections (report_id, position, title, instructions, data_sources, section_type)
                    VALUES (%s, %s, %s, %s, '[]', %s)
                    RETURNING *
                """, (report_id, start_pos + i, s.get('title', f'Section {i+1}'), s.get('instructions'), s.get('section_type', 'prose')))
                created_sections.append(_serialize(dict(cur.fetchone())))

            # Only set brief if the report doesn't already have one
            if brief_text and not report.get('report_brief'):
                cur.execute(
                    "UPDATE cvc.trend_reports SET report_brief = %s, updated_at = NOW() WHERE id = %s",
                    (brief_text, report_id)
                )
        conn.commit()

    return {"sections_created": len(created_sections), "sections": created_sections, "brief_set": bool(brief_text and not report.get('report_brief'))}


@router.post("/{report_id}/import-outline/upload")
async def import_outline_upload(
    report_id: int,
    file: UploadFile = File(...),
    user: UserInfo = Depends(require_jwt),
):
    """Upload a PDF and extract its outline via LLM."""
    _get_report_or_404(report_id)
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files supported")

    content = await file.read()
    tmp_path = f"/tmp/outline_import_{report_id}_{file.filename}"
    with open(tmp_path, 'wb') as f:
        f.write(content)

    content_text = ""
    try:
        import pdfplumber
        with pdfplumber.open(tmp_path) as pdf:
            pages = []
            for p in pdf.pages[:40]:
                t = p.extract_text()
                if t:
                    pages.append(t)
            content_text = "\n\n".join(pages)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"PDF extraction failed: {e}")
    finally:
        try:
            os.remove(tmp_path)
        except FileNotFoundError:
            pass

    if not content_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from PDF")

    # Reuse the text import endpoint logic
    from fastapi import Request
    fake_req = ImportOutlineRequest(content_text=content_text)
    return await import_outline(report_id, fake_req, user)


# ── Report Brief ──────────────────────────────────────────────────────────────

@router.post("/{report_id}/brief")
def generate_brief(report_id: int, background_tasks: BackgroundTasks,
                   user: UserInfo = Depends(require_jwt)):
    """Generate a report brief from the title, theme, and section outline."""
    report = _get_report_or_404(report_id)
    background_tasks.add_task(_generate_brief_bg, report_id)
    return {"ok": True, "status": "generating"}


def _generate_brief_bg(report_id: int):
    try:
        import sys
        sys.path.insert(0, str(_REPO_ROOT))
        from workers.trends.report_builder import build_brief
        build_brief(report_id)
    except Exception as e:
        logger.error(f"Brief generation failed [{report_id}]: {e}", exc_info=True)


# ── Publish (assemble HTML) ───────────────────────────────────────────────────

@router.post("/{report_id}/publish")
def publish_report(report_id: int, background_tasks: BackgroundTasks,
                   user: UserInfo = Depends(require_jwt)):
    """Assemble all section content into final HTML."""
    _get_report_or_404(report_id)
    background_tasks.add_task(_publish_report_bg, report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.trend_reports SET status = 'generating', updated_at = NOW() WHERE id = %s",
                (report_id,)
            )
        conn.commit()
    return {"ok": True, "status": "generating"}


def _publish_report_bg(report_id: int):
    try:
        import sys
        sys.path.insert(0, str(_REPO_ROOT))
        from workers.trends.report_builder import assemble_report
        assemble_report(report_id)
    except Exception as e:
        logger.error(f"Report publish failed [{report_id}]: {e}", exc_info=True)
        try:
            with get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE cvc.trend_reports SET status = 'error', updated_at = NOW() WHERE id = %s",
                        (report_id,)
                    )
                conn.commit()
        except Exception:
            pass


# ── Reformat (sync citation re-assembly) ──────────────────────────────────────

@router.post("/{report_id}/reformat")
def reformat_report(report_id: int, user: UserInfo = Depends(require_jwt)):
    """
    Synchronously re-assemble published HTML with current citation_style.
    No LLM calls — just re-processes existing section content.
    Returns {published_html, citation_style, status}.
    Fast (~200ms).
    """
    _get_report_or_404(report_id)
    try:
        import sys
        sys.path.insert(0, str(_REPO_ROOT))
        from workers.trends.report_builder import assemble_report
        assemble_report(report_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Reformat failed: {e}")
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT published_html, citation_style, status FROM cvc.trend_reports WHERE id = %s",
                (report_id,)
            )
            row = _serialize(dict(cur.fetchone()))
    return row


# ── Rewrite (tone/audience voice pass on all done sections) ───────────────────

@router.post("/{report_id}/rewrite")
def rewrite_report(report_id: int, background_tasks: BackgroundTasks,
                   user: UserInfo = Depends(require_jwt)):
    """
    Apply the current audience + tone settings to all done sections via the voice pass.
    Does NOT regenerate from sources — rewrites existing content through the editorial pass.
    """
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.trend_reports SET status = 'generating', updated_at = NOW() WHERE id = %s",
                (report_id,)
            )
        conn.commit()
    background_tasks.add_task(_rewrite_report_bg, report_id)
    return {"ok": True, "status": "generating"}


def _rewrite_report_bg(report_id: int):
    try:
        import sys
        sys.path.insert(0, str(_REPO_ROOT))
        from workers.trends.report_builder import rewrite_sections
        rewrite_sections(report_id)
    except Exception as e:
        logger.error(f"Report rewrite failed [{report_id}]: {e}", exc_info=True)
        try:
            with get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE cvc.trend_reports SET status = 'error', updated_at = NOW() WHERE id = %s",
                        (report_id,)
                    )
                conn.commit()
        except Exception:
            pass


# ── DOCX Download ─────────────────────────────────────────────────────────────

@router.get("/{report_id}/download/docx")
def download_docx(report_id: int, user: UserInfo = Depends(require_jwt)):
    """Generate and stream a Word document from all completed sections."""
    from fastapi.responses import StreamingResponse
    import io
    import re as _re
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM cvc.trend_reports WHERE id = %s", (report_id,))
            report = cur.fetchone()
            if not report:
                raise HTTPException(status_code=404, detail="Report not found")
            report = dict(report)
            cur.execute("""
                SELECT * FROM cvc.report_sections
                WHERE report_id = %s AND content IS NOT NULL
                ORDER BY position, id
            """, (report_id,))
            sections = [dict(r) for r in cur.fetchall()]

    doc = Document()

    # ── Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(report['title'])
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # ── Meta line
    meta_parts = [_team_name]
    if report.get('sector'):
        meta_parts.append(report['sector'])
    if report.get('date_from') and report.get('date_to'):
        meta_parts.append(f"{report['date_from']} – {report['date_to']}")
    meta = doc.add_paragraph(' · '.join(meta_parts))
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x78, 0x75, 0x69)

    # ── Report brief (italicised callout)
    if report.get('report_brief'):
        doc.add_paragraph()
        brief_para = doc.add_paragraph(report['report_brief'])
        brief_para.runs[0].italic = True
        brief_para.runs[0].font.size = Pt(10)
        brief_para.runs[0].font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()  # spacer

    # ── Sections
    for sec in sections:
        # Section heading
        heading = doc.add_paragraph()
        h_run = heading.add_run(sec['title'])
        h_run.bold = True
        h_run.font.size = Pt(14)
        h_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        content = sec.get('content') or ''
        # Strip markdown bold markers (**text**) for clean Word output
        content = _re.sub(r'\*\*(.+?)\*\*', r'\1', content)
        content = _re.sub(r'\*(.+?)\*', r'\1', content)

        # Split into paragraphs and add each
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        for para_text in paragraphs:
            # Flatten any remaining single newlines
            para_text = para_text.replace('\n', ' ')
            p = doc.add_paragraph(para_text)
            p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(8)

        doc.add_paragraph()  # spacer between sections

    # ── Stream as download
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = _re.sub(r'[^\w\s-]', '', report['title']).strip().replace(' ', '_')
    filename = f"{safe_title}_Report.docx"

    return StreamingResponse(
        buf,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


# ── Data Explorer ─────────────────────────────────────────────────────────────

_SCHEMA_DIGEST = """Available tables in the CVC database (PostgreSQL, schema: cvc):

IMPORTANT RULES FOR QUERY GENERATION:
1. stage data in cvc.companies is inconsistently cased ('Seed','seed','series_a','Series A' are all present).
   Always normalize stage with a CASE expression. Use this pattern:
   CASE
     WHEN LOWER(stage) IN ('pre-seed','pre_seed') THEN 'Pre-Seed'
     WHEN LOWER(stage) IN ('seed') THEN 'Seed'
     WHEN LOWER(stage) IN ('series a','series_a') THEN 'Series A'
     WHEN LOWER(stage) IN ('series b','series_b') THEN 'Series B'
     WHEN LOWER(stage) IN ('series c','series_c') THEN 'Series C'
     WHEN LOWER(stage) IN ('series d','series_d','series e','series_e','series f','series_f','series_h','series_i') THEN 'Series D+'
     WHEN LOWER(stage) IN ('growth') THEN 'Growth'
     ELSE 'Other/Unknown'
   END AS stage_normalized
   Then GROUP BY stage_normalized and set x_key = "stage_normalized".
2. When grouping by TWO dimensions, x_key must be the dimension shown on the axis (not the filter/sector).
   A bar chart of stage distribution should GROUP BY stage_normalized, x_key="stage_normalized", y_key="count".
3. Always alias COUNT(*) AS count, SUM(...) AS total, AVG(...) AS avg_score etc. x_key/y_key must exactly match an alias.

cvc.companies — 1,700+ tracked startups
  name (text), sector (text: 'Robotics','Supply Chain','Manufacturing','Industrial Automation','Physical AI'),
  stage (text — RAW, inconsistent casing, always normalize — see rule 1 above),
  score_composite (float 0-10), score_irs (float), score_commercial (float), score_tdf (float),
  is_portfolio (bool), founded (int — year founded e.g. 2019), country (text), hq_city (text),
  total_raised_usd (bigint), employee_count (int)
  Example stage query: SELECT CASE WHEN LOWER(stage) IN ('seed') THEN 'Seed' WHEN LOWER(stage) IN ('series a','series_a') THEN 'Series A' ELSE 'Other' END AS stage_normalized, COUNT(*) AS count FROM cvc.companies WHERE sector = 'Robotics' GROUP BY stage_normalized ORDER BY count DESC LIMIT 20

cvc.funding_rounds — funding history per company
  company_id (FK → cvc.companies.id), round_type (text), amount_usd (bigint), announced_date (date)
  Example: SELECT c.sector, SUM(fr.amount_usd) AS total FROM cvc.companies c JOIN cvc.funding_rounds fr ON fr.company_id = c.id GROUP BY c.sector ORDER BY total DESC LIMIT 10

cvc.briefing_insights — enriched intelligence from podcasts and articles (x_key: week_start or sector, y_key: count)
  id (int), week_start (date), source_type (text: 'podcast','article'), source_title (text),
  insight (text), expert (text), confidence (text: 'HIGH','MEDIUM','LOW'), sector (text), created_at (timestamptz)
  Example: SELECT TO_CHAR(week_start, 'YYYY-MM') AS month, COUNT(*) AS count FROM cvc.briefing_insights GROUP BY month ORDER BY month LIMIT 24

cvc.weekly_signals — weekly market signal summaries (one row per week)
  week_start (date), week_end (date), total_items (int), podcast_count (int), news_count (int), article_count (int),
  sentiment_positive (int), sentiment_neutral (int), sentiment_negative (int),
  top_tags (jsonb), top_companies (jsonb), top_technologies (jsonb), created_at (timestamptz)
  Example: SELECT TO_CHAR(week_start, 'YYYY-MM') AS month, SUM(total_items) AS count FROM cvc.weekly_signals GROUP BY month ORDER BY month LIMIT 24

cvc.partner_intros — introductions between startups and corporate partners
  company_id (FK → cvc.companies.id), partner_id (FK → cvc.partners.id),
  startup_name (text), partner_name (text), intro_date (date), outcome (text), status_1 (text), status_2 (text)
  Join to companies: JOIN cvc.companies c ON c.id = pi.company_id (NOT on name columns)
  Example: SELECT pi.partner_name, COUNT(*) AS count FROM cvc.partner_intros pi JOIN cvc.companies c ON c.id = pi.company_id WHERE c.sector = 'Robotics' GROUP BY pi.partner_name ORDER BY count DESC LIMIT 15

cvc.content_items — raw intelligence items (podcasts, articles, signals)
  title (text), content_type (text: 'podcast','article','signal'), url (text),
  published_at (timestamptz), sentiment (text), briefing_flag (text),
  tags (jsonb — array of strings), summary (text), created_at (timestamptz)
  Example: SELECT content_type, COUNT(*) AS count FROM cvc.content_items GROUP BY content_type ORDER BY count DESC LIMIT 10
  Note: no sector column — use tags JSONB to filter by topic, e.g. WHERE tags::text ILIKE '%robotics%'
"""


class ExploreRequest(BaseModel):
    include_section_context: bool = True


class PreviewQueryRequest(BaseModel):
    sql: str
    chart_type: str = "bar"
    x_key: str = ""
    y_key: str = ""


@router.post("/{report_id}/explore")
def explore_report(report_id: int, req: ExploreRequest, user: UserInfo = Depends(require_jwt)):
    """Call LLM to suggest 5-8 relevant data queries based on report context."""
    import requests as _req

    report = _get_report_or_404(report_id)

    # Gather section context if requested
    section_context = ""
    if req.include_section_context:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT title, instructions FROM cvc.report_sections WHERE report_id = %s ORDER BY position, id",
                    (report_id,)
                )
                secs = cur.fetchall()
        if secs:
            section_context = "\n".join(
                f"- {s['title']}: {s['instructions'] or '(no instructions)'}"
                for s in secs
            )

    _key = os.environ.get("OPENROUTER_REPORT_KEY") or os.environ.get("OPENROUTER_API_KEY", "")

    system = """You are a data analyst for a venture capital firm. Given a report's context and a database schema, suggest 5-8 relevant SQL queries that would provide useful data for the report.

Return ONLY a valid JSON array with no markdown fences. Each element must have exactly these fields:
{
  "title": "<short descriptive title>",
  "rationale": "<one sentence explaining why this data is relevant to the report>",
  "sql": "<valid PostgreSQL SELECT query using only the tables and columns described>",
  "chart_type": "<one of: bar, line, pie, area>",
  "x_key": "<column name to use for x-axis or pie labels>",
  "y_key": "<column name to use for y-axis or pie values>",
  "description": "<one sentence describing what the data shows>"
}

Rules:
- Only use tables and columns that exist in the schema provided
- All queries must be SELECT only — no INSERT, UPDATE, DELETE, DROP, etc.
- Queries must be valid PostgreSQL syntax
- Use LIMIT clauses (max 50 rows for chart data)
- Prefer aggregated/grouped data for chart visualization
- Choose chart_type based on the data: bar for comparisons, line for trends over time, pie for proportions
- x_key and y_key must exactly match column names or aliases in your SELECT clause
- Make suggestions specific to the report's sector and section topics
- Return nothing outside the JSON array"""

    user_msg = f"""Report title: {report.get('title', '')}
Sector: {report.get('sector') or 'Not specified'}
Report brief: {report.get('report_brief') or 'Not written yet'}

Section outline:
{section_context or '(no sections defined yet)'}

Database schema:
{_SCHEMA_DIGEST}

Suggest 5-8 relevant data queries for this report."""

    try:
        resp = _req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {_key}", "Content-Type": "application/json"},
            json={
                "model": "qwen/qwen3-30b-a3b",
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user_msg},
                ],
                "temperature": 0.2,
                "chat_template_kwargs": {"thinking": False},
            },
            timeout=45,
        )
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()
        # Strip thinking blocks (qwen3 extended thinking)
        raw = re.sub(r'<think>[\s\S]*?</think>', '', raw, flags=re.IGNORECASE).strip()
        # Strip markdown code fences if present
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        suggestions = json.loads(raw.strip())
        if not isinstance(suggestions, list):
            raise ValueError("LLM did not return a JSON array")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM suggestion failed: {e}")

    return {"suggestions": suggestions}


# ── Annotations (Editorial) ───────────────────────────────────────────────────

class CreateAnnotationRequest(BaseModel):
    scope: str = 'inline'              # 'inline' | 'document'
    selected_text: Optional[str] = None
    comment: str


class PatchAnnotationRequest(BaseModel):
    status: Optional[str] = None
    comment: Optional[str] = None
    proposed_rewrite: Optional[str] = None


@router.get("/{report_id}/annotations")
def list_annotations(report_id: int, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT * FROM cvc.report_annotations
                WHERE report_id = %s ORDER BY created_at
            """, (report_id,))
            return [_serialize(dict(row)) for row in cur.fetchall()]


@router.post("/{report_id}/annotations")
def create_annotation(report_id: int, req: CreateAnnotationRequest,
                      user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO cvc.report_annotations
                    (report_id, scope, selected_text, comment, status, created_by)
                VALUES (%s, %s, %s, %s, 'open', %s)
                RETURNING *
            """, (report_id, req.scope, req.selected_text, req.comment, user.username))
            row = _serialize(dict(cur.fetchone()))
        conn.commit()
    return row


@router.patch("/{report_id}/annotations/{ann_id}")
def patch_annotation(report_id: int, ann_id: int, req: PatchAnnotationRequest,
                     user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    updates: dict = {}
    if req.status is not None:
        updates['status'] = req.status
        if req.status in ('addressed', 'dismissed'):
            updates['addressed_at'] = 'NOW()'
    if req.comment is not None:
        updates['comment'] = req.comment
    if req.proposed_rewrite is not None:
        updates['proposed_rewrite'] = req.proposed_rewrite
    if not updates:
        return {"ok": True}

    # Handle NOW() specially (can't parameterise it)
    now_keys = {k for k, v in updates.items() if v == 'NOW()'}
    set_parts = []
    values = []
    for k, v in updates.items():
        if k in now_keys:
            set_parts.append(f"{k} = NOW()")
        else:
            set_parts.append(f"{k} = %s")
            values.append(v)
    values.append(ann_id)

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE cvc.report_annotations SET {', '.join(set_parts)} WHERE id = %s AND report_id = {report_id} RETURNING *",
                values
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Annotation not found")
            result = _serialize(dict(row))
        conn.commit()
    return result


@router.delete("/{report_id}/annotations/{ann_id}")
def delete_annotation(report_id: int, ann_id: int, user: UserInfo = Depends(require_jwt)):
    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM cvc.report_annotations WHERE id = %s AND report_id = %s",
                (ann_id, report_id)
            )
        conn.commit()
    return {"ok": True}


@router.post("/{report_id}/annotations/{ann_id}/address")
def address_annotation(report_id: int, ann_id: int, user: UserInfo = Depends(require_jwt)):
    """
    Call LLM to generate a proposed rewrite for this annotation.
    Inline scope: rewrites the selected_text per the comment instruction.
    Document scope: returns a general editorial note about the full report.
    Returns {proposed_rewrite: str}.
    """
    import requests as _req

    _get_report_or_404(report_id)
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM cvc.report_annotations WHERE id = %s AND report_id = %s",
                        (ann_id, report_id))
            ann = cur.fetchone()
    if not ann:
        raise HTTPException(status_code=404, detail="Annotation not found")
    ann = dict(ann)

    _key = os.environ.get("OPENROUTER_REPORT_KEY") or os.environ.get("OPENROUTER_API_KEY", "")

    if ann['scope'] == 'inline' and ann.get('selected_text'):
        system = (
            "You are a precise editorial assistant. The user will give you a text passage and an instruction. "
            "Return ONLY the revised replacement text — no explanation, no preamble, no quotes around it. "
            "Preserve the original voice and style. Match the approximate length unless the instruction specifies otherwise."
        )
        user_msg = f"Text to revise:\n{ann['selected_text']}\n\nInstruction: {ann['comment']}"
    else:
        system = (
            "You are an editorial assistant for a venture intelligence report. "
            "The user has a document-level comment about the full report. "
            "Return a concise, actionable editorial suggestion (2-4 sentences) addressing their concern. "
            "No preamble, no sign-off."
        )
        user_msg = f"Document-level editorial comment: {ann['comment']}"

    try:
        resp = _req.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {_key}", "Content-Type": "application/json"},
            json={
                "model": "moonshotai/kimi-k2.6",
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user_msg},
                ],
                "temperature": 0.3,
                "max_tokens": 1024,
            },
            timeout=45,
        )
        resp.raise_for_status()
        proposed = resp.json()["choices"][0]["message"]["content"].strip()
        import re as _re
        proposed = _re.sub(r'<think>[\s\S]*?</think>', '', proposed, flags=_re.IGNORECASE).strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM address failed: {e}")

    # Persist the proposed rewrite
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE cvc.report_annotations SET proposed_rewrite = %s WHERE id = %s RETURNING *",
                (proposed, ann_id)
            )
            row = _serialize(dict(cur.fetchone()))
        conn.commit()

    return {"proposed_rewrite": proposed, "annotation": row}


@router.post("/{report_id}/preview-query")
def preview_query(report_id: int, req: PreviewQueryRequest, user: UserInfo = Depends(require_jwt)):
    """Run a SELECT query and return up to 50 rows for preview/charting."""
    _get_report_or_404(report_id)

    sql = req.sql.strip()
    if any(kw in sql.upper() for kw in ('INSERT', 'UPDATE', 'DELETE', 'DROP', 'TRUNCATE', 'ALTER', 'CREATE', 'GRANT', 'REVOKE')):
        raise HTTPException(status_code=400, detail="Only SELECT queries are allowed")

    try:
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(sql)
                cols = [d.name for d in cur.description]
                all_rows = cur.fetchmany(50)
                rows = [_serialize(dict(row)) for row in all_rows]
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Query error: {e}")

    return {
        "columns": cols,
        "rows": rows,
        "row_count": len(rows),
        "chart_type": req.chart_type,
        "x_key": req.x_key,
        "y_key": req.y_key,
    }
