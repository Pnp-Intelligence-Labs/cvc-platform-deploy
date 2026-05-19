"""
format/agent.py — Format Bot: IC Memo PDF + Appendix PDF renderer.

Reads overview.json + appendix.json and renders two PDFs:
    1. IC Memo PDF     — clean synthesis for IC committee
    2. Appendix PDF    — full raw findings, sources, methodology

Visual identity: Plug and Play brand guidelines
    - Dark Blue:    #253B49 (headers, sidebar, dark elements)
    - Yellow:       #F0E545 (accents, horizontal rules, highlights)
    - Background:   #F5F5F7 (page / section backgrounds)
    - Body text:    #313C51
    - Secondary:    #676E7A
    - Supply Chain: #32749A (optional secondary accent — CVC sector)
    - Font:         Trebuchet MS (PnP approved digital fallback)
    - Cards:        Rounded corners (8px), drop shadows (rgba 0,0,0,0.14 / blur 14px)

Both PDFs are written to:
    workdir/[company]/[company]_IC_Memo.pdf
    workdir/[company]/[company]_Appendix.pdf

Optionally uploads both to Google Drive DD Reports folder.

Run:
    python3 -m format.agent "Dyna Robotics"
    python3 -m format.agent "Dyna Robotics" --no-upload
"""

import json
from pathlib import Path
import time
import argparse
from datetime import datetime

# ── Path setup ────────────────────────────────────────────────────────────────
from config.settings import WORKDIR, GDRIVE_CREDS, GDRIVE_TOKEN, GDRIVE_DD_FOLDER

# ── WeasyPrint ────────────────────────────────────────────────────────────────
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    print("Warning: weasyprint not installed. Install with: pip install weasyprint")


# ── Brand constants ───────────────────────────────────────────────────────────

PNP_CSS = """
@import url('data:text/css,');

/* ── Reset + base ─────────────────────────────────────────────── */
* { box-sizing: border-box; margin: 0; padding: 0; }

body {
    font-family: 'Trebuchet MS', 'Lucida Grande', Arial, sans-serif;
    font-size: 10pt;
    line-height: 1.55;
    color: #313C51;
    background: #F5F5F7;
}

/* ── Page setup ───────────────────────────────────────────────── */
@page {
    size: Letter;
    margin: 0.75in 0.7in 0.65in 0.7in;
    @bottom-right {
        content: counter(page);
        font-family: 'Trebuchet MS', Arial, sans-serif;
        font-size: 8pt;
        color: #676E7A;
    }
}
@page :first {
    @bottom-right { content: ""; }
}

/* ── Cover page ───────────────────────────────────────────────── */
.cover {
    page-break-after: always;
    background: #253B49;
    color: white;
    padding: 2.5in 1in 1.5in 1in;
    min-height: 9in;
    display: block;
}
.cover-logo {
    font-size: 11pt;
    font-weight: bold;
    letter-spacing: 0.18em;
    text-transform: uppercase;
    color: #F0E545;
    margin-bottom: 0.4in;
}
.cover-title {
    font-size: 28pt;
    font-weight: bold;
    color: white;
    line-height: 1.2;
    margin-bottom: 0.15in;
}
.cover-subtitle {
    font-size: 12pt;
    color: #B8C4CC;
    margin-bottom: 0.5in;
}
.cover-rule {
    width: 2.5in;
    height: 3px;
    background: #F0E545;
    margin-bottom: 0.4in;
}
.cover-meta {
    font-size: 9pt;
    color: #8FA5B0;
    line-height: 1.9;
}
.cover-meta span {
    color: #D0DDE3;
    font-weight: bold;
}
.cover-rec {
    display: inline-block;
    margin-top: 0.4in;
    padding: 8px 18px;
    font-size: 11pt;
    font-weight: bold;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    border-radius: 4px;
}
.rec-strong_interest { background: #2E8B57; color: white; }
.rec-proceed         { background: #32749A; color: white; }
.rec-conditional     { background: #C08000; color: white; }
.rec-pass            { background: #8B2020; color: white; }

/* ── Section headings ─────────────────────────────────────────── */
h1 {
    font-size: 15pt;
    font-weight: bold;
    color: #253B49;
    margin-top: 0.3in;
    margin-bottom: 6pt;
    page-break-after: avoid;
}
h2 {
    font-size: 11pt;
    font-weight: bold;
    color: #253B49;
    margin-top: 14pt;
    margin-bottom: 5pt;
    page-break-after: avoid;
}
h3 {
    font-size: 9.5pt;
    font-weight: bold;
    color: #313C51;
    margin-top: 10pt;
    margin-bottom: 4pt;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    page-break-after: avoid;
}

/* ── Yellow rule (PnP horizontal divider) ─────────────────────── */
.rule {
    width: 100%;
    height: 2px;
    background: #F0E545;
    margin: 10pt 0 12pt 0;
}
.rule-thin {
    width: 100%;
    height: 1px;
    background: #E0E0E5;
    margin: 8pt 0 8pt 0;
}

/* ── Cards ────────────────────────────────────────────────────── */
.card {
    background: white;
    border-radius: 8px;
    box-shadow: 0 2px 14px rgba(0,0,0,0.14);
    padding: 14pt 16pt;
    margin-bottom: 12pt;
    page-break-inside: avoid;
}
.card-blue {
    background: #253B49;
    color: white;
    border-radius: 8px;
    box-shadow: 0 2px 14px rgba(0,0,0,0.14);
    padding: 14pt 16pt;
    margin-bottom: 12pt;
    page-break-inside: avoid;
}
.card-blue h2, .card-blue h3 { color: #F0E545; }
.card-blue p, .card-blue li  { color: #D0DDE3; }

/* ── Key metrics grid ─────────────────────────────────────────── */
.metrics-grid {
    display: table;
    width: 100%;
    margin-bottom: 12pt;
    border-collapse: separate;
    border-spacing: 6pt;
}
.metric-cell {
    display: table-cell;
    width: 25%;
    background: white;
    border-radius: 8px;
    box-shadow: 0 2px 14px rgba(0,0,0,0.14);
    padding: 10pt 12pt;
    text-align: left;
    vertical-align: top;
}
.metric-label {
    font-size: 7.5pt;
    color: #676E7A;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    margin-bottom: 3pt;
}
.metric-value {
    font-size: 12pt;
    font-weight: bold;
    color: #253B49;
}

/* ── Flags ────────────────────────────────────────────────────── */
.flag-red {
    background: #FFF0F0;
    border-left: 4px solid #C0392B;
    border-radius: 0 6px 6px 0;
    padding: 8pt 10pt;
    margin-bottom: 7pt;
    page-break-inside: avoid;
}
.flag-yellow {
    background: #FFFBF0;
    border-left: 4px solid #F0E545;
    border-radius: 0 6px 6px 0;
    padding: 8pt 10pt;
    margin-bottom: 7pt;
    page-break-inside: avoid;
}
.flag-label {
    font-size: 7.5pt;
    font-weight: bold;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    color: #676E7A;
    margin-bottom: 2pt;
}
.flag-red .flag-label   { color: #C0392B; }
.flag-yellow .flag-label { color: #8A7200; }
.flag-text { font-size: 9pt; color: #313C51; }
.flag-source { font-size: 8pt; color: #676E7A; margin-top: 2pt; }

/* ── IC questions ─────────────────────────────────────────────── */
.ic-question {
    background: white;
    border-radius: 6px;
    box-shadow: 0 1px 8px rgba(0,0,0,0.10);
    padding: 9pt 11pt;
    margin-bottom: 8pt;
    page-break-inside: avoid;
}
.ic-question-high { border-left: 4px solid #253B49; }
.ic-question-medium { border-left: 4px solid #32749A; }
.ic-question-low { border-left: 4px solid #B0BEC5; }
.ic-q-text { font-size: 9.5pt; font-weight: bold; color: #253B49; }
.ic-q-meta { font-size: 7.5pt; color: #676E7A; margin-top: 2pt; }

/* ── Section summary blocks ───────────────────────────────────── */
.section-block {
    margin-bottom: 12pt;
    page-break-inside: avoid;
}
.section-label {
    font-size: 7.5pt;
    font-weight: bold;
    color: #32749A;
    text-transform: uppercase;
    letter-spacing: 0.09em;
    margin-bottom: 3pt;
}
.section-text { font-size: 9pt; color: #313C51; }

/* ── Thesis block ─────────────────────────────────────────────── */
.thesis {
    background: #F0F4F6;
    border-left: 4px solid #253B49;
    border-radius: 0 6px 6px 0;
    padding: 10pt 14pt;
    margin: 10pt 0 14pt 0;
    font-size: 9.5pt;
    color: #253B49;
    font-style: italic;
}

/* ── Cross-signal chips ───────────────────────────────────────── */
.signal {
    background: white;
    border-radius: 6px;
    box-shadow: 0 1px 8px rgba(0,0,0,0.10);
    padding: 8pt 10pt;
    margin-bottom: 6pt;
    page-break-inside: avoid;
}
.signal-red    { border-left: 4px solid #C0392B; }
.signal-yellow { border-left: 4px solid #F0E545; }
.signal-green  { border-left: 4px solid #2E8B57; }
.signal-headline { font-size: 9pt; font-weight: bold; color: #253B49; }
.signal-detail   { font-size: 8.5pt; color: #676E7A; margin-top: 2pt; }

/* ── Appendix-specific ────────────────────────────────────────── */
.agent-section-header {
    background: #253B49;
    color: white;
    border-radius: 8px 8px 0 0;
    padding: 10pt 14pt 8pt 14pt;
    page-break-after: avoid;
}
.agent-section-header h2 { color: #F0E545; margin: 0; font-size: 12pt; }
.agent-section-header p  { color: #8FA5B0; font-size: 8.5pt; margin-top: 3pt; font-style: italic; }
.agent-section-body {
    background: white;
    border-radius: 0 0 8px 8px;
    box-shadow: 0 2px 14px rgba(0,0,0,0.14);
    padding: 12pt 14pt 14pt 14pt;
    margin-bottom: 16pt;
}

.finding-row {
    border-bottom: 1px solid #EAECEE;
    padding: 7pt 0;
    page-break-inside: avoid;
}
.finding-row:last-child { border-bottom: none; }
.finding-id     { font-size: 7pt; color: #9AA5B4; font-family: monospace; }
.finding-topic  { font-size: 7.5pt; font-weight: bold; color: #253B49; text-transform: uppercase; letter-spacing: 0.06em; }
.finding-text   { font-size: 9pt; color: #313C51; margin: 2pt 0; }
.finding-claimed { font-size: 8.5pt; color: #32749A; }
.finding-delta  { font-size: 8.5pt; color: #676E7A; font-style: italic; }
.verdict-badge {
    display: inline-block;
    font-size: 7pt;
    font-weight: bold;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    padding: 1pt 5pt;
    border-radius: 3px;
    margin-right: 5pt;
}
.verdict-confirmed         { background: #EAF7EF; color: #2E8B57; }
.verdict-contradicts_claim { background: #FEECEC; color: #C0392B; }
.verdict-unverified_claim  { background: #FFF8E6; color: #8A7200; }
.verdict-no_claim          { background: #EEF2F7; color: #32749A; }
.verdict-not_found         { background: #F0F1F3; color: #676E7A; }

.source-row {
    font-size: 8pt;
    color: #676E7A;
    padding: 3pt 0;
    border-bottom: 1px solid #F0F1F3;
}
.source-row a { color: #32749A; }
.source-row:last-child { border-bottom: none; }

.meta-table {
    width: 100%;
    border-collapse: collapse;
    font-size: 8.5pt;
}
.meta-table td {
    padding: 4pt 8pt;
    border-bottom: 1px solid #EAECEE;
    color: #313C51;
}
.meta-table td:first-child { color: #676E7A; width: 45%; }

/* ── Utility ──────────────────────────────────────────────────── */
p { margin-bottom: 6pt; }
ul { padding-left: 16pt; margin-bottom: 6pt; }
li { margin-bottom: 3pt; font-size: 9pt; }
.page-break { page-break-before: always; }
.no-break   { page-break-inside: avoid; }
.muted      { color: #676E7A; }
.mono       { font-family: monospace; font-size: 8pt; }
"""


# ── HTML builders ─────────────────────────────────────────────────────────────

def _esc(text) -> str:
    """HTML-escape a string."""
    if text is None:
        return ""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _rec_label(rec: str) -> str:
    return {
        "strong_interest": "Strong Interest",
        "proceed":         "Proceed to Next Stage",
        "conditional":     "Conditional — Address Issues",
        "pass":            "Pass",
    }.get(rec, rec.replace("_", " ").title())


def build_ic_memo_html(overview: dict) -> str:
    company   = _esc(overview.get("company", ""))
    date      = _esc(overview.get("date", datetime.now().strftime("%Y-%m-%d")))
    one_liner = _esc(overview.get("one_liner", ""))
    stage     = _esc(overview.get("stage", ""))
    raise_amt = _esc(overview.get("raise_amount", ""))
    valuation = _esc(overview.get("valuation_ask", ""))
    sector    = _esc(overview.get("sector", ""))
    rec       = overview.get("recommendation", "")
    rationale = _esc(overview.get("recommendation_rationale", ""))
    thesis    = _esc(overview.get("investment_thesis", ""))
    summary   = _esc(overview.get("summary", ""))

    key_metrics       = overview.get("key_metrics", {})
    section_summaries = overview.get("section_summaries", {})
    ic_questions      = overview.get("ic_questions", [])
    all_flags         = overview.get("all_flags", [])
    cross_signals     = overview.get("cross_agent_signals", [])
    meta              = overview.get("meta", {})

    # ── Cover ────────────────────────────────────────────────────────────────
    rec_class = f"rec-{rec}" if rec else "rec-proceed"
    cover_html = f"""
    <div class="cover">
        <div class="cover-logo">Plug and Play &nbsp;·&nbsp; Due Diligence</div>
        <div class="cover-title">{company}</div>
        <div class="cover-subtitle">{one_liner}</div>
        <div class="cover-rule"></div>
        <div class="cover-meta">
            <div><span>Stage:</span> {stage}</div>
            <div><span>Raise:</span> {raise_amt}</div>
            <div><span>Valuation:</span> {valuation}</div>
            <div><span>Sector:</span> {sector}</div>
            <div><span>Date:</span> {date}</div>
        </div>
        <div class="cover-rec {rec_class}">{_rec_label(rec)}</div>
    </div>
    """

    # ── Key metrics ──────────────────────────────────────────────────────────
    metrics_html = ""
    if key_metrics and isinstance(key_metrics, dict):
        cells = ""
        for label, value in list(key_metrics.items())[:8]:
            label_clean = label.replace("_", " ").title()
            cells += f"""
            <td class="metric-cell">
                <div class="metric-label">{_esc(label_clean)}</div>
                <div class="metric-value">{_esc(str(value))}</div>
            </td>"""
        # Fill remaining cells if odd number
        remainder = len(key_metrics) % 4
        if remainder:
            for _ in range(4 - remainder):
                cells += '<td class="metric-cell" style="background:transparent;box-shadow:none;"></td>'
        metrics_html = f'<table class="metrics-grid"><tr>{cells}</tr></table>'

    # ── Investment thesis ─────────────────────────────────────────────────────
    thesis_html = f'<div class="thesis">{thesis}</div>' if thesis else ""

    # ── Section summaries ─────────────────────────────────────────────────────
    SECTION_LABELS = {
        "financials":  "Financials",
        "comp":        "Market & Competitive",
        "qualitative": "Team & Founders",
        "product":     "Product & Technology",
        "news":        "News & Press",
    }
    summaries_html = ""
    if section_summaries and isinstance(section_summaries, dict):
        for key, text in section_summaries.items():
            label = SECTION_LABELS.get(key, key.replace("_", " ").title())
            summaries_html += f"""
            <div class="section-block">
                <div class="section-label">{_esc(label)}</div>
                <div class="section-text">{_esc(str(text))}</div>
            </div>"""

    # ── Cross-agent signals ───────────────────────────────────────────────────
    signals_html = ""
    if cross_signals:
        for sig in cross_signals:
            sev       = sig.get("severity", "yellow")
            headline  = _esc(sig.get("headline", ""))
            detail    = _esc(sig.get("detail", ""))
            sig_class = f"signal-{sev}" if sev in ("red", "yellow", "green") else "signal-yellow"
            signals_html += f"""
            <div class="signal {sig_class}">
                <div class="signal-headline">{headline}</div>
                {'<div class="signal-detail">' + detail + '</div>' if detail else ''}
            </div>"""

    # ── Flags ─────────────────────────────────────────────────────────────────
    red_flags    = [f for f in all_flags if f.get("severity") == "red"]
    yellow_flags = [f for f in all_flags if f.get("severity") == "yellow"]

    def _flag_html(flags, flag_class):
        html = ""
        for f in flags:
            agent      = _esc(f.get("agent", ""))
            topic      = _esc(f.get("topic", ""))
            finding    = _esc(f.get("our_finding", ""))
            reason     = _esc(f.get("flag_reason", ""))
            verdict    = _esc(f.get("verdict", ""))
            fid        = _esc(f.get("finding_id", ""))
            label      = "Red Flag" if flag_class == "flag-red" else "Yellow Flag"
            html += f"""
            <div class="{flag_class}">
                <div class="flag-label">{label} &nbsp;·&nbsp; {agent} &nbsp;·&nbsp; {topic}</div>
                <div class="flag-text">{finding}</div>
                {'<div class="flag-source">' + reason + '</div>' if reason else ''}
                {'<div class="flag-source muted">' + verdict + ' &nbsp;·&nbsp; ' + fid + '</div>' if fid else ''}
            </div>"""
        return html

    flags_html = _flag_html(red_flags, "flag-red") + _flag_html(yellow_flags, "flag-yellow")

    # ── IC questions ──────────────────────────────────────────────────────────
    questions_html = ""
    priority_order = {"high": 0, "medium": 1, "low": 2}
    sorted_qs = sorted(ic_questions,
                       key=lambda q: priority_order.get(q.get("priority", "low"), 2))
    for q in sorted_qs:
        priority    = q.get("priority", "medium")
        question    = _esc(q.get("question", ""))
        context     = _esc(q.get("context", ""))
        agents      = ", ".join(q.get("source_agents", []))
        finding_ids = ", ".join(q.get("finding_ids", []))
        q_class     = f"ic-question ic-question-{priority}"
        questions_html += f"""
        <div class="{q_class}">
            <div class="ic-q-text">{question}</div>
            {'<div class="ic-q-meta">' + context + '</div>' if context else ''}
            <div class="ic-q-meta muted">
                {('[' + agents + ']') if agents else ''}
                {('&nbsp;·&nbsp; ' + finding_ids) if finding_ids else ''}
                &nbsp;·&nbsp; {priority.upper()} PRIORITY
            </div>
        </div>"""

    # ── Recommendation rationale ──────────────────────────────────────────────
    rationale_html = f"""
    <div class="card">
        <h3>Rationale</h3>
        <p>{rationale}</p>
        {'<p class="muted" style="font-size:8pt;">' + summary + '</p>' if summary else ''}
    </div>""" if rationale else ""

    # ── Pipeline meta footer ──────────────────────────────────────────────────
    agents_used  = ", ".join(meta.get("agents_used", []))
    n_flags      = meta.get("total_flags", len(all_flags))
    n_red        = meta.get("red_flags", len(red_flags))
    n_yellow     = meta.get("yellow_flags", len(yellow_flags))
    n_signals    = meta.get("cross_signals", len(cross_signals))
    n_qs         = len(ic_questions)

    footer_html = f"""
    <div style="margin-top: 0.4in; border-top: 1px solid #EAECEE; padding-top: 8pt;">
        <p class="muted" style="font-size: 7.5pt;">
            Generated {date} &nbsp;·&nbsp; Agents: {_esc(agents_used)}
            &nbsp;·&nbsp; {n_flags} flags ({n_red} red, {n_yellow} yellow)
            &nbsp;·&nbsp; {n_signals} cross-agent signals
            &nbsp;·&nbsp; {n_qs} IC questions
        </p>
        <p class="muted" style="font-size: 7pt; margin-top: 3pt;">
            This document is confidential and prepared for internal investment committee use only.
            See Appendix for full methodology, raw findings, and source citations.
        </p>
    </div>"""

    # ── Assemble full HTML ────────────────────────────────────────────────────
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <title>{company} — IC Memo</title>
    <style>{PNP_CSS}</style>
</head>
<body>

{cover_html}

<!-- ── Executive summary ──────────────────────────────────────── -->
<h1>Executive Summary</h1>
<div class="rule"></div>

{metrics_html}

{thesis_html}

<div class="card">
    <h3>Recommendation &nbsp;·&nbsp; <span class="{rec_class}" style="padding: 2pt 8pt; border-radius: 4px; font-size: 8pt;">{_rec_label(rec)}</span></h3>
    {rationale_html}
</div>

<!-- ── Section summaries ──────────────────────────────────────── -->
{'<h1 class="page-break">Section Summaries</h1><div class="rule"></div>' + summaries_html if summaries_html else ''}

<!-- ── Cross-agent signals ────────────────────────────────────── -->
{'<h1 class="page-break">Cross-Agent Signals</h1><div class="rule"></div>' + signals_html if signals_html else ''}

<!-- ── Flags ──────────────────────────────────────────────────── -->
{'<h1 class="page-break">Due Diligence Flags</h1><div class="rule"></div>' + flags_html if flags_html else ''}

<!-- ── IC questions ───────────────────────────────────────────── -->
{'<h1 class="page-break">IC Questions</h1><div class="rule"></div>' + questions_html if questions_html else ''}

{footer_html}

</body>
</html>"""

    return html


def build_appendix_html(appendix: dict, overview: dict) -> str:
    company = _esc(appendix.get("company", overview.get("company", "")))
    date    = _esc(appendix.get("date", datetime.now().strftime("%Y-%m-%d")))
    meta    = appendix.get("meta", {})

    VERDICT_LABELS = {
        "confirmed":         "Confirmed",
        "contradicts_claim": "Contradicts",
        "unverified_claim":  "Unverified",
        "no_claim":          "No Claim",
        "not_found":         "Not Found",
    }

    # ── Cover ─────────────────────────────────────────────────────────────────
    cover_html = f"""
    <div class="cover">
        <div class="cover-logo">Plug and Play &nbsp;·&nbsp; Due Diligence Appendix</div>
        <div class="cover-title">{company}</div>
        <div class="cover-subtitle">Raw Findings, Sources &amp; Methodology</div>
        <div class="cover-rule"></div>
        <div class="cover-meta">
            <div><span>Date:</span> {date}</div>
            <div><span>Agents:</span> {_esc(', '.join(meta.get('agents_used', [])))}</div>
            <div><span>Total findings:</span> {meta.get('total_findings', '?')}</div>
            <div><span>Total sources:</span> {meta.get('total_sources', '?')}</div>
        </div>
    </div>
    """

    # ── Per-agent sections ────────────────────────────────────────────────────
    agent_sections_html = ""
    for section in appendix.get("agent_sections", []):
        agent       = section.get("agent", "")
        title       = _esc(section.get("title", ""))
        mental      = _esc(section.get("mental_model", ""))
        summary     = _esc(section.get("summary", ""))
        findings    = section.get("findings", [])
        section_meta = section.get("meta", {})

        # Findings rows
        findings_rows = ""
        for f in findings:
            if not isinstance(f, dict):
                continue
            fid      = _esc(f.get("id", ""))
            topic    = _esc(f.get("topic", ""))
            text     = _esc(f.get("our_finding", ""))
            claimed  = f.get("claimed")
            delta    = f.get("delta")
            verdict  = f.get("verdict", "")
            conf     = _esc(f.get("confidence", ""))
            is_flag  = f.get("flag", False)

            verdict_label = VERDICT_LABELS.get(verdict, verdict)
            verdict_class = f"verdict-{verdict}" if verdict in VERDICT_LABELS else "verdict-no_claim"
            flag_mark     = ' <span style="color:#C0392B; font-weight:bold;">⚑</span>' if is_flag else ""

            sources = f.get("sources", [])
            src_html = ""
            for s in sources[:3]:
                stitle = _esc(s.get("title", s.get("url", "")))
                surl   = _esc(s.get("url", ""))
                sdate  = _esc(s.get("date", ""))
                src_html += f'<div class="source-row"><a href="{surl}">{stitle}</a> {sdate}</div>'

            claimed_html = f'<div class="finding-claimed">Claimed: {_esc(str(claimed))}</div>' if claimed else ""
            delta_html   = f'<div class="finding-delta">Delta: {_esc(str(delta))}</div>' if delta else ""

            findings_rows += f"""
            <div class="finding-row">
                <div>
                    <span class="finding-id">{fid}</span>
                    &nbsp;
                    <span class="finding-topic">{topic}</span>
                    &nbsp;
                    <span class="verdict-badge {verdict_class}">{verdict_label}</span>
                    <span class="muted" style="font-size:7pt;">{conf}</span>
                    {flag_mark}
                </div>
                <div class="finding-text">{text}</div>
                {claimed_html}
                {delta_html}
                {src_html}
            </div>"""

        # Section meta table
        meta_rows = ""
        for key, val in section_meta.items():
            meta_rows += f"<tr><td>{_esc(key.replace('_', ' ').title())}</td><td>{_esc(str(val))}</td></tr>"

        agent_sections_html += f"""
        <div class="no-break">
            <div class="agent-section-header">
                <h2>{title}</h2>
                <p>{mental}</p>
            </div>
        </div>
        <div class="agent-section-body">
            {'<p><em>' + summary + '</em></p><div class="rule-thin"></div>' if summary else ''}
            {findings_rows if findings_rows else '<p class="muted">No findings.</p>'}
            {'<div class="rule-thin"></div><table class="meta-table">' + meta_rows + '</table>' if meta_rows else ''}
        </div>"""

    # ── Sources ───────────────────────────────────────────────────────────────
    sources_html = ""
    for src in appendix.get("sources", []):
        stitle = _esc(src.get("title", ""))
        surl   = _esc(src.get("url", ""))
        sdate  = _esc(src.get("date", ""))
        agent  = _esc(src.get("agent", ""))
        fid    = _esc(src.get("finding_id", ""))
        sources_html += f"""
        <div class="source-row">
            <strong><a href="{surl}">{stitle or surl}</a></strong>
            &nbsp; <span class="muted">{sdate}</span>
            &nbsp; <span class="muted">{agent} · {fid}</span>
        </div>"""

    # ── Methodology ──────────────────────────────────────────────────────────
    methodology = appendix.get("methodology", {})

    meth_summary = f"""
    <table class="meta-table">
        <tr><td>Total docs processed</td><td>{methodology.get('total_docs_processed', '?')}</td></tr>
        <tr><td>Total web searches</td><td>{methodology.get('total_web_searches', '?')}</td></tr>
        <tr><td>Total LLM passes</td><td>{methodology.get('total_llm_passes', '?')}</td></tr>
        <tr><td>Total pipeline time</td><td>{methodology.get('total_pipeline_seconds', '?')}s</td></tr>
        <tr><td>Primary model</td><td>{_esc(methodology.get('models_used', {}).get('primary', ''))}</td></tr>
        <tr><td>Fallback model</td><td>{_esc(methodology.get('models_used', {}).get('fallback', ''))}</td></tr>
    </table>"""

    flag_rules    = methodology.get("flag_severity_rules", {})
    verdict_tax   = methodology.get("verdict_taxonomy", {})

    flag_rules_html = "".join(
        f"<tr><td style='font-weight:bold;'>{_esc(k)}</td><td>{_esc(v)}</td></tr>"
        for k, v in flag_rules.items()
    )
    verdict_html = "".join(
        f"<tr><td style='font-weight:bold;'>{_esc(k)}</td><td>{_esc(v)}</td></tr>"
        for k, v in verdict_tax.items()
    )

    meth_html = f"""
    {meth_summary}
    {'<h3 style="margin-top:12pt;">Flag Severity Rules</h3><table class="meta-table">' + flag_rules_html + '</table>' if flag_rules_html else ''}
    {'<h3 style="margin-top:12pt;">Verdict Taxonomy</h3><table class="meta-table">' + verdict_html + '</table>' if verdict_html else ''}
    """

    # ── Assemble ─────────────────────────────────────────────────────────────
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <title>{company} — Appendix</title>
    <style>{PNP_CSS}</style>
</head>
<body>

{cover_html}

<!-- ── Agent findings ─────────────────────────────────────────── -->
<h1>Findings by Agent</h1>
<div class="rule"></div>
{agent_sections_html}

<!-- ── All sources ────────────────────────────────────────────── -->
<h1 class="page-break">All Sources Cited</h1>
<div class="rule"></div>
<div class="card">
{sources_html if sources_html else '<p class="muted">No sources recorded.</p>'}
</div>

<!-- ── Methodology ────────────────────────────────────────────── -->
<h1 class="page-break">Methodology</h1>
<div class="rule"></div>
<div class="card">
{meth_html}
</div>

<div style="margin-top: 0.4in; border-top: 1px solid #EAECEE; padding-top: 8pt;">
    <p class="muted" style="font-size: 7.5pt;">
        Generated {date} &nbsp;·&nbsp; Confidential — Internal Use Only
    </p>
</div>

</body>
</html>"""

    return html


# ── Google Drive upload ───────────────────────────────────────────────────────

def upload_to_drive(file_path: Path, filename: str) -> str | None:
    """Upload a file to Google Drive DD Reports folder. Returns file URL or None."""
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaFileUpload
        import google.auth
    except ImportError:
        print("  Warning: google-auth not installed. Skipping Drive upload.")
        return None

    if not GDRIVE_CREDS.exists():
        print(f"  Warning: Google credentials not found at {GDRIVE_CREDS}. Skipping upload.")
        return None

    try:
        creds = None
        if GDRIVE_TOKEN.exists():
            creds = Credentials.from_authorized_user_file(str(GDRIVE_TOKEN))
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                print("  Warning: Google Drive token invalid. Skipping upload.")
                return None

        service = build("drive", "v3", credentials=creds)

        # Find or create DD Reports folder
        results = service.files().list(
            q=f"name='{GDRIVE_DD_FOLDER}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields="files(id, name)"
        ).execute()
        folders = results.get("files", [])

        if folders:
            folder_id = folders[0]["id"]
        else:
            folder_meta = {
                "name":     GDRIVE_DD_FOLDER,
                "mimeType": "application/vnd.google-apps.folder",
            }
            folder = service.files().create(body=folder_meta, fields="id").execute()
            folder_id = folder["id"]

        # Upload file
        file_meta = {"name": filename, "parents": [folder_id]}
        media     = MediaFileUpload(str(file_path), mimetype="application/pdf")
        uploaded  = service.files().create(
            body=file_meta, media_body=media, fields="id, webViewLink"
        ).execute()

        url = uploaded.get("webViewLink", "")
        print(f"  Uploaded: {url}")
        return url

    except Exception as e:
        print(f"  Warning: Drive upload failed: {e}")
        return None


# ── PDF rendering ─────────────────────────────────────────────────────────────

def render_pdf(html: str, output_path: Path) -> bool:
    """Render HTML to PDF using WeasyPrint. Returns True on success."""
    if not WEASYPRINT_AVAILABLE:
        print(f"  Skipping PDF render — weasyprint not available.")
        return False
    try:
        WeasyHTML(string=html).write_pdf(str(output_path))
        size_kb = output_path.stat().st_size // 1024
        print(f"  PDF written: {output_path} ({size_kb} KB)")
        return True
    except Exception as e:
        print(f"  Error rendering PDF: {e}")
        return False


# ── Main ──────────────────────────────────────────────────────────────────────

def run(company: str, upload: bool = True, version: str = None) -> dict:
    safe_name   = company.replace(" ", "_").replace("/", "-")
    company_dir = WORKDIR / safe_name
    suffix      = f"_{version}" if version else ""

    label = f"Format Bot: {company}" + (f" [{version}]" if version else "")
    print(f"\n{label}")
    print("=" * 50)

    start = time.time()

    # ── Load data ─────────────────────────────────────────────────────────────
    overview_path = company_dir / f"overview{suffix}.json"
    appendix_path = company_dir / f"appendix{suffix}.json"

    if not overview_path.exists():
        print(f"Error: {overview_path.name} not found")
        print(f"Run overview agent first: python3 -m overview.agent \"{company}\"")
        return {"status": "failed", "error": f"{overview_path.name} missing"}

    overview = json.loads(overview_path.read_text())
    appendix = json.loads(appendix_path.read_text()) if appendix_path.exists() else {}

    if not appendix:
        print(f"Warning: {appendix_path.name} not found — Appendix PDF will be minimal.")

    # ── Output paths ──────────────────────────────────────────────────────────
    memo_pdf_path     = company_dir / f"{safe_name}_IC_Memo{suffix}.pdf"
    appendix_pdf_path = company_dir / f"{safe_name}_Appendix{suffix}.pdf"

    # Also write HTML debug files
    memo_html_path     = company_dir / f"{safe_name}_IC_Memo{suffix}.html"
    appendix_html_path = company_dir / f"{safe_name}_Appendix{suffix}.html"

    # ── Build + render IC Memo ─────────────────────────────────────────────────
    print("\nBuilding IC Memo PDF...")
    memo_html = build_ic_memo_html(overview)
    memo_html_path.write_text(memo_html)
    memo_ok   = render_pdf(memo_html, memo_pdf_path)

    # ── Build + render Appendix ───────────────────────────────────────────────
    print("\nBuilding Appendix PDF...")
    appendix_html = build_appendix_html(appendix, overview)
    appendix_html_path.write_text(appendix_html)
    appendix_ok   = render_pdf(appendix_html, appendix_pdf_path)

    # ── Upload to Drive ───────────────────────────────────────────────────────
    memo_url     = None
    appendix_url = None

    if upload:
        print("\nUploading to Google Drive...")
        if memo_ok:
            memo_url = upload_to_drive(
                memo_pdf_path,
                f"{company} IC Memo {overview.get('date', '')}.pdf"
            )
        if appendix_ok:
            appendix_url = upload_to_drive(
                appendix_pdf_path,
                f"{company} Appendix {overview.get('date', '')}.pdf"
            )

    elapsed = int(time.time() - start)
    result  = {
        "company":         company,
        "date":            datetime.now().strftime("%Y-%m-%d"),
        "agent":           "format",
        "status":          "complete" if (memo_ok or appendix_ok) else "failed",
        "memo_pdf":        str(memo_pdf_path) if memo_ok else None,
        "appendix_pdf":    str(appendix_pdf_path) if appendix_ok else None,
        "memo_url":        memo_url,
        "appendix_url":    appendix_url,
        "meta": {
            "total_seconds": elapsed,
        },
    }

    print(f"\nDone ({elapsed}s)")
    if memo_ok:
        print(f"  IC Memo:  {memo_pdf_path}")
    if appendix_ok:
        print(f"  Appendix: {appendix_pdf_path}")

    return result


# ── Review Memo renderer ──────────────────────────────────────────────────────

def _build_review_banner_text(delta: dict) -> tuple[str, str]:
    """Returns (banner_html, banner_plain) for the review delta."""
    n_removed   = len(delta.get("flags_removed", []))
    n_added     = len(delta.get("flags_added", []))
    rec_changed = delta.get("recommendation_changed", False)
    summary     = delta.get("analyst_summary", "")
    rec_note_html  = (f" Recommendation changed from <strong>{delta.get('original_recommendation','—')}</strong>."
                      if rec_changed else "")
    rec_note_plain = (f" Recommendation changed from {delta.get('original_recommendation','—')}."
                      if rec_changed else "")
    banner_html = f"""
        <div style="background:#FFF3CD;border:1px solid #F59E0B;border-radius:8px;padding:12px 16px;margin-bottom:20px;">
          <div style="font-weight:700;font-size:10pt;color:#92400E;margin-bottom:4px;">
            \u2713 Analyst-Reviewed Memo
          </div>
          <div style="font-size:9pt;color:#78350F;line-height:1.5;">
            {n_removed} flag(s) removed &nbsp;\u00b7&nbsp; {n_added} flag(s) added.{rec_note_html}
            {f'<br>{summary}' if summary else ''}
          </div>
        </div>"""
    banner_plain = (
        f"ANALYST-REVIEWED MEMO\n"
        f"{n_removed} flag(s) removed · {n_added} flag(s) added.{rec_note_plain}\n"
        f"{summary}\n"
    )
    return banner_html, banner_plain


def _render_review_docx(company: str, review_memo: dict, out_path: Path) -> bool:
    """
    Render the review memo as a DOCX for manual editing.
    Produces a clean, lightly styled Word document.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        def heading(text: str, level: int = 1):
            p = doc.add_heading(text, level=level)
            p.runs[0].font.color.rgb = RGBColor(0x25, 0x3B, 0x49)

        def body(text: str):
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(10) if p.runs else None

        def kv(key: str, val: str):
            p = doc.add_paragraph()
            run_k = p.add_run(f"{key}: ")
            run_k.bold = True
            run_k.font.size = Pt(10)
            run_v = p.add_run(str(val or "—"))
            run_v.font.size = Pt(10)

        # Title
        title = doc.add_heading(f"IC Memo (Analyst-Reviewed) — {company}", 0)
        title.runs[0].font.color.rgb = RGBColor(0x25, 0x3B, 0x49)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")

        # Review delta banner
        delta = review_memo.get("review_delta", {})
        if delta:
            _, banner_plain = _build_review_banner_text(delta)
            p = doc.add_paragraph(banner_plain)
            p.runs[0].font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        # Key fields
        heading("Summary", 1)
        kv("One-liner",       review_memo.get("one_liner", "—"))
        kv("Recommendation",  review_memo.get("recommendation", "—"))
        kv("Rationale",       review_memo.get("recommendation_rationale", "—"))
        kv("Stage",           review_memo.get("stage", "—"))
        kv("Raise Amount",    review_memo.get("raise_amount", "—"))
        kv("Valuation Ask",   review_memo.get("valuation_ask", "—"))

        # Sections
        for section_key in ("thesis", "financials", "market", "team", "product", "risks", "ic_questions"):
            section_data = review_memo.get(section_key) or review_memo.get("sections", {}).get(section_key)
            if not section_data:
                continue
            heading(section_key.replace("_", " ").title(), 2)
            if isinstance(section_data, str):
                body(section_data)
            elif isinstance(section_data, dict):
                for k, v in section_data.items():
                    if isinstance(v, list):
                        kv(k.replace("_", " ").title(), "")
                        for item in v:
                            doc.add_paragraph(f"• {item}", style="List Bullet")
                    else:
                        kv(k.replace("_", " ").title(), str(v))
            elif isinstance(section_data, list):
                for item in section_data:
                    doc.add_paragraph(f"• {item}", style="List Bullet")

        # Flags
        all_flags = review_memo.get("all_flags", [])
        if all_flags:
            heading("Flags", 2)
            for f in all_flags:
                if isinstance(f, dict):
                    sev = f.get("severity", "").upper()
                    doc.add_paragraph(f"[{sev}] {f.get('description', str(f))}", style="List Bullet")

        doc.save(str(out_path))
        print(f"  DOCX written: {out_path} ({out_path.stat().st_size // 1024} KB)")
        return True
    except Exception as e:
        print(f"  Error rendering DOCX: {e}")
        return False


def render_review_memo(company: str, review_memo: dict) -> dict:
    """
    Render a corrected IC memo as both PDF and DOCX from review_memo.json.
    Returns dict with 'pdf' and 'docx' Path keys (None if render failed).
    """
    safe_name   = company.replace(" ", "_").replace("/", "-")
    company_dir = WORKDIR / safe_name
    pdf_path    = company_dir / f"{safe_name}_Review_Memo.pdf"
    docx_path   = company_dir / f"{safe_name}_Review_Memo.docx"
    html_path   = company_dir / f"{safe_name}_Review_Memo.html"

    delta = review_memo.get("review_delta", {})
    banner_html, _ = _build_review_banner_text(delta) if delta else ("", "")

    html = build_ic_memo_html(review_memo)
    html = html.replace("<body>", f"<body>{banner_html}", 1)
    html_path.write_text(html)

    pdf_ok  = render_pdf(html, pdf_path)
    docx_ok = _render_review_docx(company, review_memo, docx_path)

    return {
        "pdf":  pdf_path  if pdf_ok  else None,
        "docx": docx_path if docx_ok else None,
    }


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="DD Format Bot — IC Memo + Appendix PDF")
    parser.add_argument("company",     help="Company name (e.g. 'Dyna Robotics')")
    parser.add_argument("--no-upload", action="store_true",
                        help="Skip Google Drive upload")
    args = parser.parse_args()
    run(args.company, upload=not args.no_upload)
